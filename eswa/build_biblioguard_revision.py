# -*- coding: utf-8 -*-
"""Build the novelty-focused ESWA resubmission package.

All BiblioGuard values are injected from the released JSON/NPZ artifacts.
The package deliberately separates the main retrieval experiment from the
older generation/deployment feasibility evidence.
"""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parent
OUT = BASE.parent / "ESWA_submission_revised"
OUT.mkdir(exist_ok=True)
FIGURES = BASE / "figures"
RESULTS = BASE / "results"

BG = json.loads((RESULTS / "biblioguard_results.json").read_text())
TRANSFER = json.loads(
    (RESULTS / "biblioguard_transfer_results.json").read_text()
)
TABLES = json.loads((RESULTS / "eswa_tables.json").read_text())
DIAG = json.loads((RESULTS / "metadata_diagnostics.json").read_text())
GEN = json.loads((RESULTS / "gen_eval_summary.json").read_text())

DATASETS = ("scidocs", "scifact", "nfcorpus", "trec-covid")
DATASET_NAMES = {
    "scidocs": "SCIDOCS",
    "scifact": "SciFact",
    "nfcorpus": "NFCorpus",
    "trec-covid": "TREC-COVID",
}
TITLE = (
    "BiblioGuard: Selective bibliographic metadata intervention for "
    "multi-domain scientific retrieval-augmented generation"
)
AUTHOR = "Zichen Feng"
AFFILIATION = (
    "Department of Computer Science and Technology, Sanya University, "
    "Sanya 572022, China"
)
EMAIL = "3353854381@qq.com"
ORCID = "0009-0008-5491-3370"
REPOSITORY = "https://github.com/TwistXXF/PaperPilot-Reproduction"
OLD_REFERENCE = "ESWA-D-26-34170"


def f4(value: float) -> str:
    return f"{value:.4f}"


def pct(value: float) -> str:
    return f"{100 * value:.1f}%"


def p_text(value: float) -> str:
    if value < 0.001:
        return f"{value:.2e}"
    return f"{value:.3f}"


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:tblHeader"))


def _cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def _page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def new_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)
    _page_number(section.footer.paragraphs[0])

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.05
    for style_name, size in (("Title", 15), ("Heading 1", 12), ("Heading 2", 11)):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.keep_with_next = True
    return doc


def paragraph(
    doc: Document,
    text: str = "",
    *,
    bold: bool = False,
    italic: bool = False,
    align=None,
    size: float | None = None,
    keep: bool = False,
):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size is not None:
        run.font.size = Pt(size)
    if align is not None:
        p.alignment = align
    p.paragraph_format.keep_with_next = keep
    return p


def heading(doc: Document, text: str, level: int = 1):
    return doc.add_heading(text, level=level)


def bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text)
    return p


def table(doc: Document, headers: list[str], rows: list[list[str]], size=7.5):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.autofit = True
    header = tbl.rows[0]
    _set_repeat_table_header(header)
    for index, value in enumerate(headers):
        cell = header.cells[index]
        cell.text = str(value)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(size)
    for values in rows:
        row = tbl.add_row()
        _cant_split(row)
        for index, value in enumerate(values):
            cell = row.cells[index]
            cell.text = str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(size)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return tbl


def figure(doc: Document, filename: str, caption: str, width=6.5):
    path = FIGURES / filename
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = paragraph(doc, caption, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    cap.paragraph_format.keep_with_next = False


def _assert_artifacts() -> None:
    if set(BG["results"]) != set(DATASETS):
        raise RuntimeError("BiblioGuard results must contain all four datasets")
    if set(TRANSFER["results"]) != {"scifact", "nfcorpus"}:
        raise RuntimeError("official train-to-test results are incomplete")


def build_title_page() -> Path:
    doc = new_document()
    paragraph(doc, "Title page", bold=True, size=14)
    paragraph(doc, TITLE, bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph(doc, AUTHOR, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph(doc, AFFILIATION, align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph(doc, f"E-mail: {EMAIL}", align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph(doc, f"ORCID: {ORCID}", align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph(doc, "")
    paragraph(doc, f"Corresponding author: {AUTHOR} ({EMAIL})")
    paragraph(doc, "Declarations of interest: none.")
    paragraph(doc, "Acknowledgements: none.")
    paragraph(
        doc,
        "Data and code availability: source code, frozen metadata, and "
        f"per-query results are available at {REPOSITORY}. The revised "
        "release identifier will be inserted after archival deposit.",
    )
    output = OUT / "00_Title_Page.docx"
    doc.save(output)
    return output


def build_highlights() -> Path:
    highlights = [
        "Metadata actions keep the underlying content retrieval unchanged.",
        "A conservative gate intervenes only when local evidence is positive.",
        "The method improves a strong baseline and abstains under weak evidence.",
        "Independent train-to-test tests expose the limits of transfer.",
        "Risk–coverage results quantify benefits and query-level harms.",
    ]
    if not (3 <= len(highlights) <= 5):
        raise AssertionError("Elsevier requires 3-5 highlights")
    if any(len(item) > 85 for item in highlights):
        raise AssertionError("Each highlight must be at most 85 characters")
    doc = new_document()
    paragraph(doc, "Highlights", bold=True, size=14)
    for item in highlights:
        bullet(doc, item)
    output = OUT / "02_Highlights.docx"
    doc.save(output)
    return output


def build_cover_letter() -> Path:
    doc = new_document()
    paragraph(doc, "Cover letter", bold=True, size=14)
    paragraph(doc, "Dear Editor,")
    paragraph(
        doc,
        f"Please consider our substantially revised manuscript, “{TITLE}”, "
        "for Expert Systems with Applications. An earlier version was "
        f"submitted as {OLD_REFERENCE} and was declined at the preliminary "
        "screen because its methodological novelty was insufficient. We "
        "disclose that history so that the scale of the revision is clear.",
    )
    paragraph(
        doc,
        "This is not a textual resubmission. The earlier manuscript mainly "
        "reported a four-dataset retrieval evaluation. The new manuscript "
        "introduces BiblioGuard, a query-level selective intervention "
        "algorithm. It estimates paired gains of atomic citation or recency "
        "actions among similar historical queries and applies a "
        "Bonferroni-corrected pessimistic decision score; otherwise it "
        "retains a content-only fallback.",
    )
    paragraph(
        doc,
        "The experimental design has also been rebuilt. Metadata actions now "
        "preserve exactly the same content score, fusion weights, and top-100 "
        "candidate set as the fallback. The fallback is selected from strong "
        "dense and hybrid retrievers using training folds. We add a "
        "ContextualRouter-style local mean baseline, a global-best action, "
        "uncorrected and corrected gates, word/character and neighbourhood "
        "ablations, ten repeated fold seeds, paired bootstrap intervals, "
        "two-sided tests, risk–coverage curves, and active-query harm counts. "
        "We further add official train-to-test transfer experiments on "
        "SciFact and NFCorpus, where test labels are accessed only after all "
        "routing decisions are frozen.",
    )
    paragraph(
        doc,
        "We intentionally narrow the novelty claim. BiblioGuard is presented "
        "as a metadata-specific selective decision layer, not as a new "
        "general confidence theory or a distribution-free safety guarantee. "
        "The manuscript now uses “multi-domain” rather than “cross-domain”, "
        "reports null transfer findings, and separates earlier generation and "
        "deployment observations as feasibility evidence.",
    )
    paragraph(
        doc,
        "The manuscript is original, is not under consideration elsewhere, "
        "and all authors have approved the submission. We have no competing "
        "interests to declare.",
    )
    paragraph(doc, "Sincerely,")
    paragraph(doc, f"{AUTHOR}\n{AFFILIATION}\n{EMAIL}")
    output = OUT / "03_Cover_Letter.docx"
    doc.save(output)
    return output


def build_declarations() -> Path:
    doc = new_document()
    paragraph(doc, "Declarations", bold=True, size=14)
    heading(doc, "Competing interests", 2)
    paragraph(doc, "The author declares no competing interests.")
    heading(doc, "Funding", 2)
    paragraph(doc, "This research received no specific external funding.")
    heading(doc, "Data and code availability", 2)
    paragraph(
        doc,
        f"Code, frozen metadata, and exact per-query artifacts: {REPOSITORY}. "
        "A version-specific archival identifier will be added before final "
        "submission after the revised branch is released.",
    )
    heading(doc, "Use of generative artificial intelligence", 2)
    paragraph(
        doc,
        "Generative AI tools were used for language editing and code review. "
        "The author designed the study, verified every numerical result, "
        "reviewed the final text, and accepts responsibility for the work.",
    )
    output = OUT / "04_Declarations.docx"
    doc.save(output)
    return output


def build_manuscript() -> Path:
    _assert_artifacts()
    doc = new_document()
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(TITLE)

    heading(doc, "Abstract", 1)
    sc = BG["results"]["scidocs"]
    ci = sc["paired_bootstrap_95ci"]
    paragraph(
        doc,
        "Bibliographic metadata can improve scientific retrieval, but fixed "
        "citation or recency priors can also displace relevant evidence. We "
        "formulate metadata use as a selective intervention problem and "
        "introduce BiblioGuard, a lightweight decision layer for scientific "
        "retrieval-augmented generation. For each query, BiblioGuard retrieves "
        "similar historical queries, estimates paired NDCG@10 effects for "
        "nine atomic metadata actions, and intervenes only when the best "
        "Bonferroni-corrected pessimistic decision score is positive. Every "
        "action preserves the selected content retriever's scores, fusion "
        "rule, and candidate set, removing a confound in earlier metadata "
        f"comparisons. Across four scientific benchmarks, the method raises "
        f"SCIDOCS from {f4(sc['fallback_N@10'])} to "
        f"{f4(sc['biblioguard_N@10'])} (gain {f4(sc['gain_N@10'])}; "
        f"paired-bootstrap 95% interval [{f4(ci[0])}, {f4(ci[1])}]) against "
        "a strong SBERT-Dense fallback, while conservatively abstaining on "
        "the other domains. Ten fold seeds give a mean gain of "
        f"{f4(sc['repeated_seeds']['gain_mean'])} ± "
        f"{f4(sc['repeated_seeds']['gain_std'])}. Official train-to-test "
        "evaluations on SciFact and NFCorpus show near-complete abstention and "
        "no mean improvement, exposing limited transfer rather than hiding it. "
        "Risk–coverage analysis shows that conservative gating reduces harmed "
        "queries relative to unconstrained local or global routing. The "
        "contribution is a metadata-specific selective intervention policy, "
        "not a formal distribution-free safety guarantee.",
    )
    paragraph(
        doc,
        "Keywords: scientific information retrieval; retrieval-augmented "
        "generation; bibliographic metadata; selective prediction; query "
        "routing; abstention",
        italic=True,
    )

    heading(doc, "1. Introduction", 1)
    paragraph(
        doc,
        "Scientific retrieval-augmented generation (RAG) systems must select "
        "evidence from collections in which lexical match, semantic relevance, "
        "citation impact, publication time, and venue quality are only partly "
        "aligned. Content encoders directly address topical relevance. "
        "Bibliographic signals provide a different kind of evidence: they can "
        "favour influential or recent work, but they can also encode age, "
        "field, and popularity biases. A fixed metadata prior therefore "
        "creates a decision problem rather than a universally beneficial "
        "ranking feature.",
    )
    paragraph(
        doc,
        "The central question is not whether citations are correlated with "
        "relevance on average. It is whether a system should alter a strong "
        "content ranking for the current query. This distinction matters in "
        "deployed writing assistants: an unhelpful intervention changes the "
        "evidence exposed to the generator, whereas abstention retains a "
        "known content-only ranking.",
    )
    paragraph(
        doc,
        "We address this problem with BiblioGuard. The method treats citation "
        "and recency weights as atomic actions relative to the same underlying "
        "content retriever. It uses cross-fitted query similarity to estimate "
        "local paired effects, penalises uncertain actions and multiplicity, "
        "and otherwise selects the content fallback. The design borrows the "
        "useful non-parametric intuition of retrieval-based routers but changes "
        "the target from absolute model performance to the incremental effect "
        "of a metadata intervention.",
    )
    paragraph(doc, "The paper makes four bounded contributions:")
    bullet(
        doc,
        "A confound-free action construction in which metadata is the only "
        "ranking component that changes relative to a strong content fallback.",
    )
    bullet(
        doc,
        "BiblioGuard, a metadata-specific local paired-effect policy with a "
        "multiplicity-corrected pessimistic gate and explicit abstention.",
    )
    bullet(
        doc,
        "A four-domain evaluation with local/global routing baselines, gate "
        "and representation ablations, risk–coverage analysis, query-level "
        "harm counts, paired bootstrap intervals, and repeated folds.",
    )
    bullet(
        doc,
        "An official train-to-test evaluation on two domains that documents "
        "both conservative transfer and the method's limited coverage.",
    )

    heading(doc, "2. Related work and novelty boundary", 1)
    heading(doc, "2.1. Scientific retrieval and bibliographic signals", 2)
    paragraph(
        doc,
        "Sparse and dense retrieval capture complementary matching signals, "
        "and hybrid fusion is a standard strategy on heterogeneous corpora "
        "(Thakur et al., 2021; Reimers and Gurevych, 2019; Xiao et al., 2023). "
        "Scientific search additionally exposes citation, publication-year, "
        "and venue metadata. Prior systems use metadata for filtering, "
        "reranking, source reliability, or citation-aware synthesis (Hwang et "
        "al., 2025; Yousuf et al., 2026; Ding et al., 2026). These studies "
        "motivate metadata-aware retrieval but do not imply that a fixed prior "
        "is safe for every query or corpus.",
    )
    heading(doc, "2.2. Query routing and selective prediction", 2)
    paragraph(
        doc,
        "Adaptive RAG increasingly routes queries among retrievers or evidence "
        "sources. R3AG learns query-specific retriever capability from "
        "retrieval and generation supervision (Zhao et al., 2026), while "
        "RouteRAG jointly learns graph/text retrieval decisions (Guo et al., "
        "2026). ContextualRouter estimates candidate performance from similar "
        "historical queries and reports that a simple k-nearest-neighbour mean "
        "is highly competitive (Varangot-Reille et al., 2026). We therefore "
        "include that simple local-mean rule as a direct experimental baseline.",
    )
    paragraph(
        doc,
        "Abstention and conservative policy improvement offer a second "
        "foundation. Selective prediction studies risk as coverage changes "
        "(Geifman and El-Yaniv, 2019), and high-confidence policy improvement "
        "compares a proposed policy with a baseline (Thomas et al., 2015). "
        "Recent conformal routing provides distribution-free violation-rate "
        "guarantees under explicit calibration assumptions (Uddin and Bauer, "
        "2026). BiblioGuard does not provide that guarantee. Its weighted "
        "neighbours are dependent and similarity-selected; its score must be "
        "read as an operational pessimistic decision score; it is not a "
        "formal confidence bound.",
    )
    heading(doc, "2.3. What is new here", 2)
    table(
        doc,
        ["Method family", "Prediction target", "Uncertainty/abstention", "Our distinction"],
        [
            ["ContextualRouter-style kNN", "Absolute candidate performance", "Optional cost rule", "Paired metadata effect, same-content fallback"],
            ["Safe policy improvement", "Policy return vs baseline", "Formal bounds under assumptions", "Offline retrieval action; no formal safety claim"],
            ["Selective prediction", "Prediction correctness/risk", "Reject option", "Intervention coverage and retrieval harm"],
            ["BiblioGuard", "Incremental NDCG@10 effect", "Pessimistic score; fallback", "Bibliographic intervention for scientific retrieval"],
        ],
        size=7,
    )
    paragraph(
        doc,
        "The novelty is consequently application-methodological: the action "
        "definition, paired-effect target, multiplicity-aware gate, and "
        "same-content fallback are integrated for bibliographic intervention. "
        "We do not claim a new general statistical theory or use “first”.",
    )

    heading(doc, "3. Problem formulation", 1)
    paragraph(
        doc,
        "Let q be a query, D a corpus, s₀(q,d) the score produced by a "
        "content-only retriever, and zₘ(d) a normalised metadata signal "
        "m ∈ {citation, recency}. The fallback ranking uses s₀. An atomic "
        "action a=(m,λ) reranks only the top-100 fallback candidates with "
        "sₐ(q,d) = s₀(q,d) + λzₘ(d). The action does not change the "
        "content retriever, its sparse/"
        "dense fusion, or its candidate set. Citation actions use λ ∈ "
        "{.05,.10,.15,.20,.30}; recency actions use λ ∈ "
        "{.05,.10,.15,.20}.",
    )
    paragraph(
        doc,
        "For historical query i, the paired action effect is Δᵢ(a) = "
        "NDCG@10ᵢ(a) − NDCG@10ᵢ(0), where 0 denotes fallback. The policy "
        "observes query text and historical paired "
        "effects, but not the relevance judgements of the query being routed. "
        "Its output is either one atomic action or fallback.",
    )

    heading(doc, "4. BiblioGuard", 1)
    heading(doc, "4.1. Training-fold content fallback", 2)
    paragraph(
        doc,
        "Within each outer fold, BiblioGuard selects the content-only candidate "
        "with the highest training-fold mean NDCG@10 among the candidates for "
        "which same-content actions are available. Candidates include "
        "SBERT-Dense, BGE-Dense, Neural-Hybrid, and BGE-Hybrid where document "
        "scores are available. The held-out fold does not influence this "
        "selection. This directly answers why the method should not be compared "
        "only with an arbitrary 0.5/0.5 hybrid.",
    )
    heading(doc, "4.2. Local paired-effect estimation", 2)
    paragraph(
        doc,
        "Word TF-IDF (1–2 grams) and character TF-IDF (3–5 grams) are fitted "
        "only on training queries. For a held-out query, cosine similarity "
        "selects k = ⌈√nₜʳₐᶢₙ⌉ neighbours. Non-negative similarities "
        "plus 0.001 are normalised to weights wᵢ. For every action, the "
        "weighted mean effect is μ̂ₐ = ∑ᵢ wᵢΔᵢ(a). We use the unbiased "
        "weighted variance vₐ = ∑ᵢ wᵢ(Δᵢ(a) − μ̂ₐ)² / "
        "(1 − ∑ᵢ wᵢ²), effective sample size nₑ = 1 / ∑ᵢ wᵢ², "
        "and SEₐ = √(vₐ ∑ᵢ wᵢ²).",
    )
    heading(doc, "4.3. Pessimistic decision score and abstention", 2)
    paragraph(
        doc,
        "For nine actions and family α=.05, the decision score is "
        "PDSₐ = μ̂ₐ − t(1 − α/9, ⌊nₑ⌋ − 1)SEₐ. BiblioGuard selects the "
        "action with maximal PDS only if that maximum is positive; otherwise "
        "it returns the content fallback. The Bonferroni term addresses the "
        "within-query action search. Because retrieved neighbours are "
        "similarity-selected rather than independent observations, PDS is not "
        "labelled a confidence bound.",
    )
    paragraph(doc, "Algorithm 1. Cross-fitted BiblioGuard", bold=True, keep=True)
    table(
        doc,
        ["Step", "Operation"],
        [
            ["1", "Split queries; select the strongest content fallback on outer-training labels."],
            ["2", "Construct nine same-content atomic metadata outcomes for training and held-out queries."],
            ["3", "Fit word/character TF-IDF on training text; retrieve k similar training queries."],
            ["4", "Estimate weighted paired effects, variance, effective sample size, and PDS per action."],
            ["5", "Apply argmax action if max PDS>0; otherwise return the content-only fallback."],
        ],
        size=7.5,
    )
    heading(doc, "4.4. Computational complexity", 2)
    paragraph(
        doc,
        "After content retrieval, an intervention reranks at most 100 "
        "documents. Routing requires one sparse query projection and a "
        "similarity search over historical query vectors. The reference code "
        "uses exact cosine search; approximate nearest-neighbour indexing can "
        "replace it without changing the decision rule.",
    )

    heading(doc, "5. Experimental design", 1)
    heading(doc, "5.1. Datasets and frozen metadata", 2)
    dataset_rows = []
    for dataset in DATASETS:
        row = TABLES["datasets"][dataset]
        coverage = row["meta_cov"]
        dataset_rows.append(
            [
                DATASET_NAMES[dataset],
                row["domain"],
                f"{row['n_docs']:,}",
                f"{row['n_queries']:,}",
                pct(coverage["citations"]),
                pct(coverage["year"]),
            ]
        )
    table(
        doc,
        ["Dataset", "Domain", "Documents", "Test queries", "Citation coverage", "Year coverage"],
        dataset_rows,
    )
    paragraph(
        doc,
        "SCIDOCS covers computer science; SciFact contains biomedical claims; "
        "NFCorpus covers nutrition and medicine; TREC-COVID is a temporal "
        "biomedical search collection. Citation/year/venue metadata come from "
        "the frozen repository snapshot collected from Semantic Scholar and "
        "OpenAlex during 2025–2026. No live API values are used at evaluation "
        "time. The recency reference year remains 2024 to reproduce the frozen "
        "experimental protocol.",
    )
    heading(doc, "5.2. Main and independent protocols", 2)
    paragraph(
        doc,
        "The main protocol is five-fold cross-fitting with seed 42. Held-out "
        "qrels never enter content-base selection, TF-IDF fitting, neighbour "
        "retrieval, effect estimation, or action selection. We repeat the "
        "complete procedure for seeds 0–9. For independent validation, official "
        "SciFact and NFCorpus train queries determine the base, representation, "
        "neighbours, and actions; test qrels are opened only after decisions "
        "are frozen. We call this train-to-test transfer, not external-dataset "
        "validation.",
    )
    heading(doc, "5.3. Baselines and ablations", 2)
    paragraph(
        doc,
        "Content candidates are SBERT-Dense, BGE-Dense, and their 0.5 BM25 "
        "hybrids. Routing baselines are: (i) the global training-fold best "
        "action; (ii) a ContextualRouter-style kNN local mean without an "
        "uncertainty penalty; (iii) an uncorrected one-sided t penalty; "
        "(iv) a bounded-effect empirical Bernstein penalty; and (v) "
        "BiblioGuard. The legacy PAV feature router is retained as a "
        "historical baseline. Mechanism ablations compare word-only, "
        "character-only, half-k, and double-k settings. Hyperparameters are "
        "not selected on test outcomes.",
    )
    heading(doc, "5.4. Statistical analysis", 2)
    paragraph(
        doc,
        "The primary comparison uses paired, two-sided Wilcoxon signed-rank "
        "tests, with Holm correction across four datasets. Mean paired gains "
        "receive 10,000-resample query-level bootstrap 95% intervals (seed "
        "2026). We report improved, unchanged, and harmed queries both overall "
        "and among active interventions. A penalty multiplier from 0 to 1.5 "
        "produces the risk–coverage analysis. Exact per-query arrays are "
        "released.",
    )

    heading(doc, "6. Results", 1)
    heading(doc, "6.1. Primary comparison with strong fallbacks", 2)
    primary_rows = []
    for dataset in DATASETS:
        row = BG["results"][dataset]
        ci = row["paired_bootstrap_95ci"]
        base_counts = ", ".join(
            f"{name}:{count}" for name, count in row["fallback_counts"].items()
        )
        primary_rows.append(
            [
                DATASET_NAMES[dataset],
                f4(row["fallback_N@10"]),
                f4(row["biblioguard_N@10"]),
                f"{row['gain_N@10']:+.4f}",
                f"[{ci[0]:.4f}, {ci[1]:.4f}]",
                pct(row["selection_rate"]),
                p_text(row["wilcoxon_p_holm_two_sided"]),
                base_counts,
            ]
        )
    paragraph(
        doc,
        "Table 2. Cross-fitted BiblioGuard against a training-fold-selected "
        "content fallback. CI is the paired-query bootstrap interval; p is "
        "two-sided and Holm-adjusted.",
        italic=True,
        keep=True,
    )
    table(
        doc,
        ["Dataset", "Fallback", "BiblioGuard", "Gain", "95% CI", "Active", "Holm p", "Fold bases (queries)"],
        primary_rows,
        size=6.6,
    )
    paragraph(
        doc,
        f"On SCIDOCS, every outer fold selects SBERT-Dense, and BiblioGuard "
        f"improves mean NDCG@10 by {sc['gain_N@10']:.4f}. This result exceeds "
        "the strongest content-only score rather than only a weaker hybrid. "
        "The policy abstains on the other domains under the corrected gate. "
        "The result therefore supports selective utility on one domain, not "
        "universal metadata benefit.",
    )
    figure(
        doc,
        "Fig2_biblioguard_main.png",
        "Fig. 1. Same-content interventions compared with strong fallbacks and routing baselines.",
        width=6.55,
    )

    heading(doc, "6.2. What the gate changes", 2)
    comparison_rows = []
    for dataset in DATASETS:
        row = BG["results"][dataset]
        for method, label in (
            ("global_best", "Global-best"),
            ("local_mean", "Local kNN mean"),
            ("uncorrected", "Uncorrected PDS"),
            ("empirical_bernstein", "Empirical Bernstein"),
            ("biblioguard", "BiblioGuard"),
        ):
            comp = row["comparisons"][method]
            comparison_rows.append(
                [
                    DATASET_NAMES[dataset],
                    label,
                    f4(comp["N@10"]),
                    f"{comp['gain_N@10']:+.4f}",
                    pct(comp["selection_rate"]),
                    str(comp["outcomes_active"]["improved"]),
                    str(comp["outcomes_active"]["unchanged"]),
                    str(comp["outcomes_active"]["harmed"]),
                ]
            )
    table(
        doc,
        ["Dataset", "Policy", "NDCG@10", "Gain", "Active", "Improved", "Unchanged", "Harmed"],
        comparison_rows,
        size=6.8,
    )
    sc_global = sc["comparisons"]["global_best"]
    sc_local = sc["comparisons"]["local_mean"]
    sc_bg = sc["comparisons"]["biblioguard"]
    paragraph(
        doc,
        f"The SCIDOCS global-best action obtains a larger average gain "
        f"({sc_global['gain_N@10']:+.4f}) but harms "
        f"{sc_global['outcomes_active']['harmed']} active queries. The local "
        f"mean harms {sc_local['outcomes_active']['harmed']}; BiblioGuard "
        f"reduces that count to {sc_bg['outcomes_active']['harmed']} while "
        "intervening on fewer queries. This is the central trade-off: the gate "
        "sacrifices some mean upside for lower exposure to query-level harm. "
        "It does not eliminate harm. The empirical Bernstein comparator "
        "abstains throughout the evaluated domains, illustrating "
        "that a bounded-effect penalty can be too conservative at these local "
        "sample sizes.",
    )
    figure(
        doc,
        "Fig4_active_outcomes.png",
        "Fig. 2. Improved, unchanged, and harmed queries among active BiblioGuard decisions.",
        width=6.3,
    )

    heading(doc, "6.3. Risk–coverage and repeated folds", 2)
    risk_rows = []
    for row in sc["risk_coverage"]:
        risk_rows.append(
            [
                f"{row['penalty_scale']:.2f}",
                pct(row["coverage"]),
                f"{row['gain_N@10']:+.4f}",
                str(row["harmed_active"]),
            ]
        )
    table(
        doc,
        ["Penalty multiplier", "Coverage", "Mean gain", "Harmed active queries"],
        risk_rows,
        size=7.5,
    )
    figure(
        doc,
        "Fig3_risk_coverage.png",
        "Fig. 3. SCIDOCS intervention coverage and gain as the pessimism penalty changes.",
        width=5.8,
    )
    repeated_rows = []
    for dataset in DATASETS:
        repeated = BG["results"][dataset]["repeated_seeds"]
        repeated_rows.append(
            [
                DATASET_NAMES[dataset],
                f"{repeated['score_mean']:.4f} ± {repeated['score_std']:.4f}",
                f"{repeated['gain_mean']:+.4f} ± {repeated['gain_std']:.4f}",
                f"{100 * repeated['selection_mean']:.1f}% ± {100 * repeated['selection_std']:.1f}%",
            ]
        )
    table(
        doc,
        ["Dataset", "BiblioGuard NDCG@10", "Gain", "Active rate"],
        repeated_rows,
    )
    paragraph(
        doc,
        "SCIDOCS remains positive across all ten repeated seeds. The other "
        "datasets remain conservative, showing that the null activation is not "
        "a peculiarity of seed 42.",
    )

    heading(doc, "6.4. Feature and neighbourhood ablations", 2)
    ablation_rows = [
        ["Default word+char, k", f4(sc["biblioguard_N@10"]), f"{sc['gain_N@10']:+.4f}", pct(sc["selection_rate"])],
    ]
    for key, label in (
        ("word_only", "Word only"),
        ("char_only", "Character only"),
        ("k_half", "Half k"),
        ("k_double", "Double k"),
    ):
        row = sc["ablations"][key]
        ablation_rows.append(
            [label, f4(row["N@10"]), f"{row['gain_N@10']:+.4f}", pct(row["selection_rate"])]
        )
    table(doc, ["SCIDOCS setting", "NDCG@10", "Gain", "Active"], ablation_rows)
    paragraph(
        doc,
        "Character-only and double-k variants are stronger on SCIDOCS, but "
        "they are reported as post-specified ablations and are not promoted to "
        "the primary method. Choosing them after viewing these test outcomes "
        "would inflate the result. Future work should tune representation and "
        "coverage on a separate development collection.",
    )

    heading(doc, "6.5. Official train-to-test transfer", 2)
    transfer_rows = []
    for dataset in ("scifact", "nfcorpus"):
        row = TRANSFER["results"][dataset]
        comp = row["comparisons"]["biblioguard"]
        transfer_rows.append(
            [
                DATASET_NAMES[dataset],
                str(row["n_train"]),
                str(row["n_test"]),
                row["selected_content_base"],
                f4(row["fallback_N@10"]),
                f4(row["biblioguard_N@10"]),
                pct(row["selection_rate"]),
                f"{comp['outcomes_active']['improved']}/{comp['outcomes_active']['unchanged']}/{comp['outcomes_active']['harmed']}",
            ]
        )
    table(
        doc,
        ["Dataset", "Train", "Test", "Base", "Fallback", "BiblioGuard", "Active", "I/U/H active"],
        transfer_rows,
        size=7,
    )
    sf_transfer = TRANSFER["results"]["scifact"]
    nf_transfer = TRANSFER["results"]["nfcorpus"]
    paragraph(
        doc,
        f"The strict gate activates on {pct(sf_transfer['selection_rate'])} of "
        f"SciFact and {pct(nf_transfer['selection_rate'])} of NFCorpus test "
        "queries, with zero mean gain in both. The uncorrected gate activates "
        "more often, and the local mean can transfer negatively on NFCorpus. "
        "These findings support abstention under weak evidence but also show "
        "that the learned intervention does not yet generalise broadly.",
    )

    heading(doc, "6.6. Legacy PAV routing baseline", 2)
    pav_rows = []
    for dataset in DATASETS:
        row = TABLES["router"][dataset]
        pav_rows.append(
            [
                DATASET_NAMES[dataset],
                f4(row["best_single_N@10"]),
                f4(row["routed_system"]["N@10"]),
                f"{row['macro_f1']:.3f}",
                f"{row['kappa']:.3f}",
            ]
        )
    table(doc, ["Dataset", "Best fixed", "PAV routed", "Macro-F1", "Kappa"], pav_rows)
    paragraph(
        doc,
        "The legacy surface-feature classifier does not reliably beat its best "
        "fixed method and has near-zero agreement beyond class imbalance. "
        "BiblioGuard differs by estimating paired intervention effects rather "
        "than predicting an absolute winning retriever label.",
    )

    heading(doc, "6.7. Metadata association diagnostics", 2)
    diag_rows = []
    for dataset in DATASETS:
        row = DIAG[dataset]
        ci = row["cit_rel_auc_95ci"]
        diag_rows.append(
            [
                DATASET_NAMES[dataset],
                str(row["rel_docs"]),
                str(row["nonrel_docs"]),
                f"{row['cit_rel_auc']:.3f} [{ci[0]:.3f}, {ci[1]:.3f}]",
                f"{row['rel_median_citations']:.0f}/{row['nonrel_median_citations']:.0f}",
                row["metadata_snapshot"],
            ]
        )
    table(
        doc,
        ["Dataset", "Positive docs", "Background docs", "Citation AUC (95% CI)", "Median cites P/B", "Snapshot"],
        diag_rows,
        size=6.5,
    )
    paragraph(
        doc,
        "The diagnostic is document-level and associative. Positive and "
        "background IDs are deduplicated; every dataset uses the same seeded "
        "background rule (up to 10:1, capped at 25,000; seed 7). SCIDOCS has a "
        "strong citation–relevance association, whereas NFCorpus and "
        "TREC-COVID are near chance. This pattern is consistent with, but does "
        "not causally explain, where citation intervention is useful. The AUC "
        "interval uses the stated Hanley–McNeil large-sample approximation.",
    )

    heading(doc, "7. Discussion", 1)
    heading(doc, "7.1. Practical interpretation", 2)
    paragraph(
        doc,
        "BiblioGuard should be interpreted as a selective decision layer, not "
        "as evidence that conservative routing maximises average NDCG. On "
        "SCIDOCS, the global action is stronger in mean but exposes more "
        "queries to harm. A deployment can choose a penalty multiplier based "
        "on its tolerance for coverage and query-level regressions. The default "
        "is deliberately conservative.",
    )
    paragraph(
        doc,
        "The train-to-test results are equally important. Near-total abstention "
        "protects the content baseline, but a system that almost never acts has "
        "limited added value. This failure boundary motivates better calibrated "
        "representations, domain-conditioned action spaces, and development "
        "collections with sufficient intervention signal.",
    )
    heading(doc, "7.2. Relation to end-to-end RAG", 2)
    paragraph(
        doc,
        f"An earlier feasibility experiment generated {GEN['n_paired']} paired "
        "answers with two fixed MiniLM retrieval backends. It found no "
        "significant relevance, faithfulness, or citation-precision difference "
        "between them. That experiment predates BiblioGuard and does not test "
        "the new policy. It is therefore not used as evidence of an end-to-end "
        "generation gain and is retained only in the repository as secondary "
        "feasibility material. The same qualification applies to the earlier "
        "small deployment pilot.",
    )
    heading(doc, "7.3. Limitations", 2)
    paragraph(
        doc,
        "First, the main gain occurs on one of four domains. Second, official "
        "train-to-test transfer yields abstention rather than improvement. "
        "Third, PDS is not a calibrated or distribution-free confidence bound. "
        "Fourth, the metadata snapshot is time-dependent, incomplete, and may "
        "encode popularity and age bias. Fifth, SCIDOCS representation and k "
        "ablations show meaningful sensitivity; they require independent "
        "tuning. Sixth, NDCG@10 is a retrieval surrogate and the paper does not "
        "establish a BiblioGuard generation gain. Seventh, candidate retrieval "
        "models are compact encoders rather than current large rerankers. "
        "Eighth, query-level improvements do not prevent individual harms.",
    )
    heading(doc, "7.4. Responsible use", 2)
    paragraph(
        doc,
        "Citation counts should not be treated as scientific quality. They are "
        "affected by field size, publication age, language, venue, and social "
        "visibility. The fallback and abstention mechanism reduce unnecessary "
        "use but do not remove these biases. Interfaces should expose why a "
        "metadata action was used and allow users to disable it.",
    )

    heading(doc, "8. Conclusion", 1)
    paragraph(
        doc,
        "We introduced BiblioGuard, a selective bibliographic intervention "
        "algorithm that compares atomic citation and recency actions with an "
        "unchanged strong content ranking. The method combines local paired "
        "effects, a multiplicity-corrected pessimistic score, and explicit "
        "fallback. It produces a stable improvement over SBERT-Dense on "
        "SCIDOCS while reducing harmed-query exposure relative to aggressive "
        "routing; on the remaining domains and official transfer tests it "
        "largely abstains. The negative transfer evidence sets a clear research "
        "agenda: independent calibration, broader development collections, and "
        "direct generation evaluation are required before bibliographic "
        "intervention can be considered broadly reliable.",
    )

    heading(doc, "References", 1)
    references = [
        "Asai, A., Wu, Z., Wang, Y., Sil, A., & Hajishirzi, H. (2024). Self-RAG: Learning to retrieve, generate, and critique through self-reflection. ICLR.",
        "Ding, H., Zhao, Y., Hu, T., Wang, Z., Patwardhan, M., & Cohan, A. (2026). SciRAG: Adaptive, citation-aware, and outline-guided retrieval and synthesis for scientific literature. EACL, 6440–6460.",
        "Geifman, Y., & El-Yaniv, R. (2019). SelectiveNet: A deep neural network with an integrated reject option. ICML, 2151–2159.",
        "Guo, Y., Su, M., Guan, S., Sun, Z., Jin, X., Guo, J., & Cheng, X. (2026). RouteRAG: Efficient retrieval-augmented generation from text and graph via reinforcement learning. Findings of ACL, 30042–30059. https://doi.org/10.18653/v1/2026.findings-acl.1502",
        "Hwang, J., Park, J., Park, H., Kim, D., Park, S., & Ok, J. (2025). Retrieval-augmented generation with estimation of source reliability. EMNLP, 34279–34303.",
        "Priem, J., Piwowar, H., & Orr, R. (2022). OpenAlex: A fully-open index of scholarly works, authors, venues, institutions, and concepts. arXiv:2205.01833.",
        "Reimers, N., & Gurevych, I. (2019). Sentence-BERT: Sentence embeddings using Siamese BERT-networks. EMNLP-IJCNLP, 3982–3992.",
        "Thakur, N., Reimers, N., Rücklé, A., Srivastava, A., & Gurevych, I. (2021). BEIR: A heterogeneous benchmark for zero-shot evaluation of information retrieval models. NeurIPS Datasets and Benchmarks.",
        "Thomas, P. S., Theocharous, G., & Ghavamzadeh, M. (2015). High confidence policy improvement. ICML, 2380–2388.",
        "Uddin, I., & Bauer, A. (2026). Conformal LLM routing with distribution-free safety guarantees. ACL Student Research Workshop, 791–799. https://doi.org/10.18653/v1/2026.acl-srw.70",
        "Varangot-Reille, C., Bouvard, C., & Gourru, A. (2026). Generalising LLM routing using past performance retrieval: A few-shot router is sufficient. EACL Student Research Workshop, 304–319. https://doi.org/10.18653/v1/2026.eacl-srw.22",
        "Xiao, S., Liu, Z., Zhang, P., & Muennighoff, N. (2023). C-Pack: Packaged resources to advance general Chinese embedding. arXiv:2309.07597.",
        "Yousuf, R. B., Xu, S., Sharma, M., Neeser, A., Latimer, C., & Ramakrishnan, N. (2026). Utilizing metadata for better retrieval-augmented generation. ECIR, 305–319.",
        "Zhao, T., Zhu, Y., Tian, Y., & Dou, Z. (2026). R3AG: Retriever routing for retrieval-augmented generation. ACL, 20506–20522. https://doi.org/10.18653/v1/2026.acl-long.939",
    ]
    for reference in references:
        p = paragraph(doc, reference, size=9)
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.first_line_indent = Inches(-0.2)

    heading(doc, "Data and code availability", 1)
    paragraph(
        doc,
        f"The full pipeline, frozen metadata, per-query outcomes, action "
        f"archives, and verification scripts are available at {REPOSITORY}. "
        "The exact revised commit and archival DOI will be inserted after the "
        "revision is released; no DOI from an earlier manuscript version is "
        "claimed here.",
    )
    output = OUT / "01_Manuscript_ESWA.docx"
    doc.save(output)
    return output


def main() -> None:
    paths = [
        build_title_page(),
        build_manuscript(),
        build_highlights(),
        build_cover_letter(),
        build_declarations(),
    ]
    for path in paths:
        print("saved", path)


if __name__ == "__main__":
    main()
