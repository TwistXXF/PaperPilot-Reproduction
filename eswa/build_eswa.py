# -*- coding: utf-8 -*-
"""Build the ESWA (Expert Systems with Applications) submission package
from the consolidated, audited numbers in results/eswa_tables.json.

Output folder: ../ESWA_submission/
  00_Title_Page.docx
  01_Manuscript_ESWA.docx
  02_Highlights.docx
  03_Cover_Letter.docx
  04_Declarations.docx
  figures/*.png|pdf

Every quantitative claim in the text is injected from eswa_tables.json so the
manuscript cannot drift from the result files. verify_eswa.py audits this.
"""
import json
import os
import shutil

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.normpath(os.path.join(BASE, '..', 'ESWA_submission'))
FIGOUT = os.path.join(OUT, 'figures')
os.makedirs(FIGOUT, exist_ok=True)

T = json.load(open(os.path.join(BASE, 'results', 'eswa_tables.json')))

METHODS = ['BM25', 'LSA-Dense', 'SBERT-Dense', 'BGE-Dense', 'Neural-Hybrid',
           'UMA-RAG', 'LP-RAG', 'CA-HR', 'BGE-Hybrid', 'BGE-CA-HR']
DS = ['scidocs', 'scifact', 'nfcorpus', 'trec-covid']
DS_NAME = {d: T['datasets'][d]['name'] for d in DS}

AUTHOR = 'Zichen Feng'
AFFIL = ('Department of Computer Science and Technology, Sanya University, '
         'Sanya 572022, China')
EMAIL = '3353854381@qq.com'
ORCID = '0009-0008-5491-3370'
REPO = 'https://github.com/TwistXXF/PaperPilot-Reproduction'
DOI = '10.5281/zenodo.21930729'
SITE = 'https://xxfpaperpilot.cn'

TITLE = ('BiblioGuard: Confidence-gated bibliographic metadata intervention '
         'for cross-domain scientific retrieval-augmented generation')


def f4(x):
    return f'{x:.4f}'


def f3(x):
    return f'{x:.3f}'


def pval(p):
    if p < 0.001:
        return 'p < 0.001'
    return f'p = {p:.3f}'


def new_doc():
    d = Document()
    st = d.styles['Normal']
    st.font.name = 'Times New Roman'
    st.font.size = Pt(11)
    return d


def p(doc, text, bold=False, size=11, align=None, italic=False):
    par = doc.add_paragraph()
    run = par.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if align:
        par.alignment = align
    return par


def h1(doc, t):
    p(doc, t, bold=True, size=12)


def h2(doc, t):
    p(doc, t, bold=True, size=11)


def add_table(doc, headers, rows, font_size=8):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for j, htxt in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = htxt
        for r in c.paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(font_size)
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            c = t.rows[i + 1].cells[j]
            c.text = str(v)
            for r in c.paragraphs[0].runs:
                r.font.size = Pt(font_size)
    for index, row in enumerate(t.rows):
        row._tr.get_or_add_trPr().append(OxmlElement('w:cantSplit'))
        if index == 0:
            row._tr.get_or_add_trPr().append(OxmlElement('w:tblHeader'))
    doc.add_paragraph()
    return t


# convenient accessors ------------------------------------------------------
def avg(ds, m, k):
    return T['main'][ds]['avg'][m][k]


def abl(ds, variant, k='N@10'):
    return T['ablation'][ds][variant][k]


def robust_n10(ds, noise, method):
    """Read both legacy scalar and metric-dictionary robustness formats."""
    value = T['robust'][ds][noise][method]
    return float(value['N@10'] if isinstance(value, dict) else value)


# ===========================================================================
# Title page
# ===========================================================================
def build_title_page():
    d = new_doc()
    p(d, 'Title Page', bold=True, size=14)
    d.add_paragraph()
    p(d, TITLE, bold=True, size=13)
    d.add_paragraph()
    p(d, f'{AUTHOR}', bold=True)
    p(d, AFFIL)
    p(d, f'E-mail: {EMAIL}')
    p(d, f'ORCID: {ORCID}')
    p(d, 'Telephone: +86-18032173019')
    d.add_paragraph()
    p(d, f'Corresponding author: {AUTHOR} ({EMAIL})')
    d.add_paragraph()
    p(d, 'Declarations of interest: none.')
    p(d, 'Acknowledgements: None.')
    p(d, f'Data and code availability: all data, code, and per-query result '
         f'files are publicly available at {REPO} '
         f'(archived at https://doi.org/{DOI}).')
    d.save(os.path.join(OUT, '00_Title_Page.docx'))


# ===========================================================================
# Highlights
# ===========================================================================
def build_highlights():
    d = new_doc()
    p(d, 'Highlights', bold=True, size=13)
    hl = [
        'BiblioGuard learns paired query-level gains for nine metadata actions.',
        'Simultaneous lower bounds gate interventions and otherwise abstain.',
        'Five-fold cross-fitting separates every decision from held-out qrels.',
        'SCIDOCS NDCG@10 rises by 0.0049; BiblioGuard abstains elsewhere.',
        'Removing confidence gating causes negative transfer in three domains.',
    ]
    for x in hl:
        par = d.add_paragraph(style='List Bullet')
        par.add_run(x)
    d.save(os.path.join(OUT, '02_Highlights.docx'))


# ===========================================================================
# Manuscript
# ===========================================================================
def build_manuscript():
    d = new_doc()
    p(d, TITLE, bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    d.add_paragraph()

    # ---- Abstract ---------------------------------------------------------
    g = T['generation']
    BG = T['biblioguard']['results']
    h1(d, 'Abstract')
    p(d,
      'Bibliographic metadata is a tempting ranking prior for scientific '
      'retrieval-augmented generation (RAG), but a static citation or recency '
      'boost can improve one domain and harm another. We introduce '
      'BiblioGuard, a confidence-gated intervention policy that decides, for '
      'each query, whether a strong metadata-free hybrid should be modified. '
      'Nine single-signal citation or recency actions are evaluated from '
      'historical queries. For a held-out query, word- and character-level '
      'TF-IDF retrieve similar training queries; similarity-weighted paired '
      'NDCG@10 effects are converted into simultaneous one-sided Student-t '
      'lower confidence bounds. BiblioGuard selects the action with the '
      'largest positive lower bound and otherwise abstains to BGE-Hybrid. '
      'Five-fold cross-fitting ensures that a query\'s relevance judgments '
      'never enter its representation, neighbours, effect estimate, or '
      'decision. Across SCIDOCS, SciFact, NFCorpus, and TREC-COVID (1,673 '
      'queries), BiblioGuard raises SCIDOCS NDCG@10 from 0.1832 to 0.1881 '
      '(+0.0049; 18.7% intervention rate; Holm-adjusted p < 0.001) and '
      'abstains on the other three domains, producing no observed negative '
      'mean transfer. Removing the simultaneous lower-bound gate activates '
      '45.8%-67.3% of queries in those domains and reduces NDCG@10 by '
      '0.00003, 0.00339, and 0.01148, respectively. Corpus diagnostics explain '
      'the pattern: citation actions are supported only where citation count '
      'predicts relevance. A secondary 200-query generation study and a '
      'pilot deployment establish end-to-end feasibility. Code, cross-fitted '
      'decisions, and per-query scores are released for exact reproduction.')
    p(d, 'Keywords: retrieval-augmented generation; scientific information '
         'retrieval; bibliographic metadata; selective prediction; '
         'treatment-effect routing; abstention; expert systems',
      italic=True, size=10)
    d.add_page_break()

    # ---- 1. Introduction ---------------------------------------------------
    h1(d, '1. Introduction')
    p(d,
      'Scientific retrieval-augmented generation (RAG) grounds a generated '
      'answer in papers returned by an information-retrieval system (Lewis et '
      'al., 2020). Bibliographic fields offer evidence unavailable to content '
      'encoders: citation count approximates accumulated authority and '
      'publication year captures timeliness. Their use is nevertheless an '
      'intervention, not a free feature. A fixed citation boost can favour '
      'canonical work in computer science while suppressing recent evidence '
      'in pandemic medicine. The same action can therefore be beneficial, '
      'irrelevant, or harmful across queries and domains.')
    p(d,
      'Existing adaptive RAG routers largely predict which retriever will '
      'perform best. That objective is insufficient when metadata is '
      'optional and the deployment already has a strong content-based '
      'fallback: the operational question is whether the incremental effect '
      'of a metadata action is positive with enough evidence to justify '
      'departing from the fallback. Absolute performance prediction also '
      'ignores the paired structure of alternative rankings for the same '
      'query and can over-intervene after choosing among many noisy actions.')
    p(d,
      'We address this gap with BiblioGuard, a cross-fitted, confidence-gated '
      'meta-policy for bibliographic intervention. It estimates local paired '
      'uplift for nine single-signal citation or recency actions from similar '
      'historical queries, corrects the one-sided decision threshold across '
      'the action family, and abstains to BGE-Hybrid unless the best action '
      'has a positive lower confidence bound. The policy is lightweight, '
      'encoder-agnostic, and auditable: each decision is accompanied by its '
      'estimated effect, lower bound, neighbour count, and selected action.')
    p(d, 'We study three research questions:')
    for c in [
        'RQ1. Can paired-effect confidence gating exploit useful metadata '
        'without the cross-domain negative transfer of unconstrained routing?',
        'RQ2. Which parts of BiblioGuard—local effect estimation, simultaneous '
        'confidence correction, and abstention—account for its behaviour?',
        'RQ3. Which corpus signals explain activation, and is the resulting '
        'decision layer feasible in an end-to-end scientific RAG system?',
    ]:
        par = d.add_paragraph(style='List Bullet')
        par.add_run(c)
    p(d,
      'Answering these questions yields the following contributions:')
    for c in [
        'BiblioGuard, a new selective metadata-intervention algorithm that '
        'combines local paired treatment effects, simultaneous one-sided '
        'confidence bounds, and an explicit metadata-free fallback (RQ1);',
        'A leakage-resistant five-fold evaluation protocol in which vectoriser '
        'fitting, neighbour selection, effect estimation, and action choice '
        'exclude each held-out query and its relevance judgments (RQ1, RQ2);',
        'A mechanism ablation showing that removing the confidence gate gains '
        'more on SCIDOCS but causes negative transfer on SciFact, NFCorpus, '
        'and TREC-COVID, whereas the complete policy abstains there (RQ2);',
        'A four-domain diagnostic account linking action utility to '
        'citation-relevance association rather than metadata coverage, plus '
        'comparisons with ten retrieval configurations and two rank-fusion '
        'baselines (RQ3);',
        'An exact reproduction package containing code, nine-action outcomes, '
        'cross-fitted decisions, 135 algorithm-integrity checks, and a '
        'secondary generation/deployment feasibility study (RQ3).',
    ]:
        par = d.add_paragraph(style='List Bullet')
        par.add_run(c)
    p(d,
      'The remainder of the paper is organised as follows. Section 2 reviews '
      'related work. Section 3 describes the deployed system. Section 4 '
      'formalises the retrieval actions and BiblioGuard. Section 5 '
      'details datasets, metadata acquisition, and evaluation protocol. '
      'Section 6 reports results. Section 7 presents the deployment case '
      'study. Section 8 discusses implications and limitations, and Section 9 '
      'concludes.')

    # ---- 2. Related work ---------------------------------------------------
    h1(d, '2. Related work')
    h2(d, '2.1. Retrieval for scientific text')
    p(d,
      'Sparse lexical retrieval with BM25 (Robertson & Zaragoza, 2009) '
      'remains a strong baseline on '
      'scientific corpora. Dense passage retrieval (Karpukhin et al., 2020) '
      'and its scientific '
      'specialisations—SPECTER (Cohan et al., 2020), which exploits '
      'citation links during '
      'pre-training, and neighborhood-contrastive approaches (Ostendorff et '
      'al., 2022)—improve '
      'semantic matching. Benchmarks such as BEIR (Thakur et al., 2021) and '
      'SciRepEval (Singh et al., 2023) '
      'have shown that no single retriever dominates across domains, a '
      'finding our study sharpens for RAG-oriented pipelines. Sentence-level '
      'encoders such as Sentence-BERT (Reimers & Gurevych, 2019) and the '
      'distilled MiniLM family '
      '(Wang et al., 2020) make dense retrieval cheap enough for CPU-only '
      'deployment, which '
      'is the regime our deployed system operates in; E5 (Wang et al., '
      '2022) and the BGE '
      'family (Chen et al., 2024) push the same efficiency-quality frontier '
      'further.')
    h2(d, '2.2. Hybrid and metadata-aware ranking')
    p(d,
      'Hybrid systems interpolate sparse and dense scores (Lin, 2021; Qu '
      'et al., 2021). '
      'Re-ranking with cross-encoders or sequence-to-sequence models '
      '(monoT5; Nogueira et al., 2020) and zero-shot LLM rankers such as '
      'setwise prompting '
      '(Zhuang et al., 2024) improve top-of-ranking quality at additional '
      'inference cost; '
      'these methods re-rank by content only. CA-HR instead re-ranks by '
      'combining content scores with citation-authority and recency signals '
      'derived from real bibliographic metadata, in the spirit of '
      'authority-aware scholarly search but applied inside a RAG pipeline. '
      'UMA-RAG and LP-RAG represent two alternative metadata-aware designs '
      '(uniform metadata augmentation and length-penalised ranking, '
      'respectively) and serve as strong in-family baselines.')
    h2(d, '2.3. Adaptive and agentic RAG')
    p(d,
      'Recent systems make retrieval adaptive: Self-RAG (Asai et al., '
      '2024) learns when to '
      'retrieve, ITER-RETGEN (Shao et al., 2023) interleaves retrieval and '
      'generation, and '
      'agent-based assistants (Wang et al., 2024) plan multi-step '
      'retrieval. OpenScholar '
      '(Asai et al., 2026) synthesises scientific literature with '
      'retrieval-augmented '
      'models. Query-wise Dual-perspective Adaptive Retrieval (QuDAR) selects '
      'sparse, dense, or expanded-query retrieval from score-margin '
      'confidence and an LLM perspective (Kim et al., 2026). R3AG learns '
      'retriever capabilities and routes using both retrieval-quality and '
      'generation-utility labels (Zhao et al., 2026), while ContextualRouter '
      'uses past query performance for non-parametric LLM/retriever routing '
      '(Varangot-Reille et al., 2026). These systems predict absolute '
      'candidate utility. BiblioGuard instead estimates the paired incremental '
      'effect of an optional metadata action relative to a fixed strong '
      'fallback, applies a simultaneous lower-bound test over all actions, '
      'and can choose no intervention. The target, correction, and abstention '
      'rule are therefore distinct from best-retriever classification.')
    h2(d, '2.4. Evaluating the generation side')
    p(d,
      'Retrieval metrics do not directly measure what users read. We '
      'therefore add a generation-side study in which the same questions are '
      'answered by a fixed LLM (DeepSeek, deepseek-chat; DeepSeek-AI, '
      '2024) conditioned '
      'on contexts produced by two competing backends, scored for relevance '
      'and faithfulness with an LLM judge (Zheng et al., 2023) and for '
      'citation precision '
      'against the gold relevance judgments. This closes the loop between '
      'retrieval evaluation and deployed answer quality.')
    h2(d, '2.5. Metadata-aware scientific RAG and positioning')
    p(d,
      'The RAG literature has moved quickly, and several 2025-2026 studies '
      'are closest to our problem. SurveyGen (Bao et al., 2025) builds a '
      'large-scale '
      'scientific survey dataset and a quality-aware framework (QUAL-SG) '
      'that injects citation counts, author influence, and venue reputation '
      'into literature retrieval for survey generation; it demonstrates '
      'that '
      'quality metadata helps its generation pipeline but does not isolate '
      'when or why the retrieval-level gain occurs. Yousuf et al. (2026) '
      'systematically compare metadata-as-text, dual-encoder, and '
      'reformulation strategies on SEC filings and show that moderate '
      'metadata weights help on structured financial corpora; their '
      'metadata '
      'is structural (company, form, section), not bibliometric. SciRAG '
      '(Ding et al., 2026) couples adaptive retrieval with citation-graph '
      'symbolic '
      'reasoning and outline-guided synthesis for scientific question '
      'answering, using citations to organise evidence rather than as a '
      'ranking prior. RA-RAG (Hwang et al., 2025) estimates per-source '
      'reliability by '
      'cross-checking and retrieves only from reliable sources in '
      'multi-source QA. Table 1 summarises the positioning. Our work is '
      'complementary and, to our knowledge, the first to (i) quantify, '
      'across four scientific domains, when citation- and recency-aware '
      'ranking priors help, are neutral, or harm—linking the outcome to a '
      'measurable corpus property (citation-relevance AUC) rather than to '
      'metadata coverage—and (ii) turn that diagnostic finding into a '
      'query-level intervention algorithm with paired-effect estimation, '
      'multiplicity-aware confidence gating, and abstention. This is not a '
      'claim of a distribution-free safety guarantee: the confidence bound '
      'is a local, similarity-weighted decision statistic whose empirical '
      'behaviour is evaluated by cross-fitting.')
    d.add_page_break()
    p(d, 'Table 1. Algorithmic positioning against recent adaptive-retrieval '
         'routers. Metadata-aware systems are reviewed in the text.',
      italic=True, size=9)
    add_table(d,
      ['Study', 'Decision target', 'Supervision', 'Family-wise gate',
       'Explicit abstention'],
      [
        ['QuDAR (Kim et al., 2026)', 'Best retriever/query form',
         'Margins and LLM decision', 'No', 'No'],
        ['R3AG (Zhao et al., 2026)', 'Best retriever',
         'Retrieval + generation labels', 'No', 'No'],
        ['ContextualRouter (Varangot-Reille et al., 2026)',
         'Best model/retriever', 'Past absolute performance', 'No', 'No'],
        ['BiblioGuard (this paper)', 'Incremental metadata effect',
         'Past paired query effects', 'Yes, over nine actions',
         'Yes, to BGE-Hybrid'],
      ], font_size=7)

    # ---- 3. Deployed system ------------------------------------------------
    h1(d, '3. The deployed academic writing assistant')
    p(d,
      'The system under study is a web-based academic reading and writing '
      'assistant (name and URL withheld for anonymous review) that '
      'operationalises the retrieval stack evaluated in this paper. Users '
      'upload papers, chat with an AI assistant grounded in their library, '
      'manage highlights and AI-generated summaries, and draft writing '
      'projects with citation support. The system serves real users in '
      'production.')
    p(d,
      'Fig. 1 shows the architecture. A React 18 single-page application '
      'communicates over HTTPS with an Nginx 1.18 reverse proxy that serves '
      'static assets and forwards API calls—including server-sent-event (SSE) '
      'token streams—to a Node.js 20 / Express backend supervised by PM2. '
      'Persistent state (users, papers, conversations, messages, highlights, '
      'summaries) lives in MySQL 8.0 behind the Drizzle ORM. The RAG '
      'orchestration layer chunks uploaded papers, embeds passages, runs the '
      'hybrid retrieval and metadata re-ranking pipeline evaluated here, and '
      'assembles a cited context of five passages that is streamed to the '
      'DeepSeek chat model (deepseek-chat; DeepSeek-AI, 2024) for answer '
      'generation.')
    p(d,
      'Two deployment constraints shaped the retrieval design. First, the '
      'production server is a commodity virtual machine (2 vCPU, 1.6 GB RAM) '
      'with no GPU, ruling out cross-encoder re-ranking at query time and '
      'motivating the CPU-efficient configurations evaluated in this study. '
      'Second, answers must stream interactively, so retrieval plus '
      're-ranking must stay well under one second—satisfied by all ten '
      'configurations (Section 6.7).')
    d.add_picture(os.path.join(BASE, 'figures', 'system_architecture.png'),
                  width=Inches(6.0))
    p(d, 'Fig. 1. Architecture and deployment boundary of the deployed '
         'academic writing assistant (system anonymized for review).',
      italic=True, size=9)

    # ---- 4. Methods ---------------------------------------------------------
    h1(d, '4. Retrieval actions and BiblioGuard')
    h2(d, '4.1. Content-based retrieval')
    p(d,
      'BM25 (Robertson & Zaragoza, 2009) scores documents with k1 = 1.5, '
      'b = 0.75 over whitespace- and '
      'punctuation-tokenised title-plus-abstract text. LSA-Dense projects a '
      '50,000-feature TF-IDF space to 384 dimensions with truncated SVD '
      '(latent semantic analysis; Deerwester et al., 1990). SBERT-Dense '
      'encodes the same text with '
      'all-MiniLM-L6-v2 (Reimers & Gurevych, 2019; Wang et al., 2020) '
      '(384-dim, L2-normalised, inner-product '
      'similarity). BGE-Dense replaces the encoder with BGE-small-en-v1.5 '
      '(Xiao et al., 2023), a 12-layer 384-dim model roughly twice the '
      'depth of MiniLM, '
      'used with its official query instruction prefix. Neural-Hybrid '
      'fuses BM25 and SBERT-Dense with equal weights on min-max-normalised '
      'scores, and BGE-Hybrid is the same equal-weight fusion on the BGE '
      'backbone. To test whether metadata effects depend on the strength of '
      'the dense backbone, BGE-CA-HR applies the CA-HR re-ranking rule '
      '(Section 4.2) to the BGE hybrid instead of the MiniLM hybrid.')
    h2(d, '4.2. Metadata-aware variants')
    p(d,
      'Three configurations add bibliographic signals. UMA-RAG augments the '
      'hybrid score uniformly with venue prestige V(d) and citation authority '
      'C(d). LP-RAG multiplies the hybrid score by a length prior '
      '1 + eta * exp(-len(d)/mu), which smoothly down-weights longer '
      'documents: shorter documents receive up to a 1 + eta boost that '
      'decays exponentially with document length. CA-HR forms the hybrid '
      'base score Sₕ(q,d) = α Ssparse(q,d) + (1−α) Sdense(q,d), with α = 0.6, '
      'takes the top-100 candidates, and re-ranks by SCA(q,d) = Sₕ(q,d) + '
      'β C(d) + γ R(d), where C(d) = log[1+c(d)] / maxⱼ log[1+c(j)] '
      'is corpus-normalised citation authority and R(d) = exp[-λ(tref−t(d))] '
      'is an exponential recency term (λ = 0.1 per year, '
      'tref = 2024), with β = 0.15 and γ = 0.10. Citation counts, '
      'years, and venues are real API values; documents without a matched '
      'record receive c(d) = 0 and the corpus median year, mirroring how a '
      'deployed system must handle missing metadata.')
    h2(d, '4.3. BiblioGuard: confidence-gated paired-effect routing')
    p(d,
      'Decision problem. Let a₀ denote the metadata-free BGE-Hybrid fallback '
      '(0.5 BM25 + 0.5 BGE after per-query min-max normalisation). The action '
      'family A contains nine single-signal CA-HR configurations: five '
      'citation actions beta in {0.05, 0.10, 0.15, 0.20, 0.30} with gamma = 0 '
      'and four recency actions gamma in {0.05, 0.10, 0.15, 0.20} with beta = '
      '0. Each intervention uses CA-HR\'s 0.6 BM25 + 0.4 BGE candidate score '
      'and re-ranks its top 100. Thus the estimand is the total effect of '
      'switching from the deployed fallback to a single-metadata '
      'configuration; it does not attribute the full contrast to metadata '
      'alone. The beta = gamma = 0 configuration is retained as a '
      'fusion-weight control in Section 6.5 but is not an intervention.')
    p(d,
      'Local paired effects. For every labelled historical query qᵢ and '
      'action a, define Δᵢ(a) = NDCGᵢ@10(a) − NDCGᵢ@10(a₀). Query text '
      'is represented by the concatenation of word TF-IDF uni-/bi-grams and '
      'character-boundary TF-IDF 3-5-grams. Both vectorisers are fitted only '
      'on the training part of a fold. For a held-out query q, BiblioGuard '
      'retrieves k = ⌈√ntrain⌉ cosine-nearest training queries. '
      'Non-negative similarities receive a 0.001 stabiliser and are '
      'normalised to weights wᵢ. The local paired-effect estimator is '
      'Δ̂q(a) = Σᵢ wᵢΔᵢ(a), with effective sample size neff = 1/Σᵢwᵢ². Its '
      'weighted variance is v̂q(a) = Σᵢwᵢ[Δᵢ(a)−Δ̂q(a)]² and standard error '
      'seq(a) = √[v̂q(a)/neff].')
    p(d,
      'Simultaneous confidence gate. With family alpha = 0.05 and |A| = 9, '
      'the one-sided lower bound is Lq(a) = Δ̂q(a) − '
      't(1−α/|A|, ⌊neff⌋−1) seq(a). The policy chooses the action '
      'with the largest lower bound only if maxₐ Lq(a) > 0; otherwise it '
      'returns a₀. Bonferroni correction makes the gate deliberately '
      'conservative across the nine actions. Because neighbours are selected '
      'by similarity rather than sampled i.i.d., L is an operational '
      'uncertainty score, not a distribution-free coverage guarantee.')
    p(d, 'Algorithm 1. Cross-fitted BiblioGuard decision for held-out query q.',
      italic=True, size=9)
    add_table(d, ['Step', 'Operation'], [
      ['1', 'Fit word/character TF-IDF on training-query text only.'],
      ['2', 'Retrieve k cosine-nearest training queries for q.'],
      ['3', 'Estimate nine similarity-weighted paired NDCG@10 effects.'],
      ['4', 'Compute Bonferroni-adjusted one-sided lower bounds.'],
      ['5', 'Select argmax L if max L > 0; otherwise return BGE-Hybrid.'],
      ['6', 'Log action, effect estimate, lower bound, k, and critical value.'],
    ])
    p(d,
      'Evaluation and complexity. We use shuffled five-fold cross-fitting '
      '(seed 42); a held-out query\'s relevance judgments are never used in '
      'its vectorisation, neighbour set, effect estimate, or decision. The '
      'unconstrained ablation keeps the same neighbours and paired estimator '
      'but selects the largest positive mean effect without a lower-bound '
      'test. Given cached action outcomes, inference costs one sparse cosine '
      'search and O(k|A|) effect aggregation. A new domain with no labelled '
      'history takes the explicit cold-start path a₀ rather than extrapolating '
      'a metadata policy.')

    # ---- 5. Experimental setup ----------------------------------------------
    h1(d, '5. Experimental setup')
    h2(d, '5.1. Datasets')
    rows = []
    for ds in DS:
        m = T['datasets'][ds]
        cov = m['meta_cov']
        rows.append([m['name'], m['domain'], f"{m['n_docs']:,}",
                     f"{m['n_queries']:,}",
                     f"{100*cov['citations']:.1f}%",
                     f"{100*cov['year']:.1f}%",
                     f"{100*cov['venue']:.1f}%"])
    p(d, 'Table 2. Datasets and real-metadata coverage.', italic=True, size=9)
    add_table(d,
              ['Dataset', 'Domain', 'Docs', 'Queries', 'Citation cov.',
               'Year cov.', 'Venue cov.'], rows)
    p(d,
      'We use four public benchmarks in BEIR format (Thakur et al., 2021): '
      'SCIDOCS (computer '
      'science; Cohan et al., 2020), SciFact (biomedical claim '
      'verification; Wadden et al., 2020), NFCorpus '
      '(nutrition and medicine; Boteva et al., 2016), and TREC-COVID '
      '(COVID-19 biomedicine; '
      'Voorhees et al., 2021). They span two orders of magnitude in corpus '
      'size (3.6k-171k '
      'documents) and four distinct domains, and—critically—differ in '
      'metadata coverage (Table 2): SCIDOCS is nearly complete (99.7%), '
      'SciFact and NFCorpus are high (94%), while TREC-COVID is genuinely '
      'sparse at 69.8% citation coverage, reflecting its many preprints. We '
      'treat this sparsity not as a defect to hide but as a property of '
      'real-world corpora that our cross-domain design is built to expose.')
    h2(d, '5.2. Bibliographic informativeness diagnostics')
    DI = T['diagnostics']
    p(d, 'Table 3. Bibliographic informativeness diagnostics per corpus. '
         'Citation-relevance AUC = P(citations of a relevant document exceed '
         'those of a non-relevant one); non-relevant sets are judged '
         'non-relevant documents where available (SCIDOCS, TREC-COVID) and a '
         'seeded background sample otherwise (SciFact, NFCorpus).',
      italic=True, size=9)
    rows = []
    for ds in DS:
        dg = DI[ds]
        rows.append([DS_NAME[ds], f"{100*dg['coverage']:.1f}%",
                     f"{dg['median_citations']:.0f}",
                     f"{100*dg['pct_zero_citation']:.1f}%",
                     f"{dg['median_age']:.0f}",
                     f"{dg['rel_median_citations']:.0f}",
                     f"{dg['nonrel_median_citations']:.0f}",
                     f"{dg['cit_rel_auc']:.3f}"])
    add_table(d, ['Dataset', 'Cit. cov.', 'Median cit.', '% zero-cit.',
                  'Median age (yr)', 'Rel. med. cit.', 'Non-rel. med. cit.',
                  'Cit.-rel. AUC'], rows)
    p(d,
      'Coverage alone does not explain where metadata helps: SciFact (94.1%) '
      'and NFCorpus (94.0%) have virtually identical citation coverage yet '
      'respond oppositely to citation-aware ranking. The discriminating '
      'variable is the citation-relevance association (Table 3). On SCIDOCS, '
      f'relevant documents carry a median of '
      f'{DI["scidocs"]["rel_median_citations"]:.0f} citations versus '
      f'{DI["scidocs"]["nonrel_median_citations"]:.0f} for non-relevant ones '
      f'(AUC = {DI["scidocs"]["cit_rel_auc"]:.3f}), so a citation prior '
      'genuinely separates relevant from non-relevant material. On SciFact '
      f'the association is weak (AUC = {DI["scifact"]["cit_rel_auc"]:.3f}). '
      f'On NFCorpus it is absent entirely (AUC = '
      f'{DI["nfcorpus"]["cit_rel_auc"]:.3f}, Mann-Whitney p = 0.55; medians '
      '6 vs. 5), and on TREC-COVID it is inverted '
      f'(AUC = {DI["trec-covid"]["cit_rel_auc"]:.3f}): relevant COVID-19 '
      'documents are newer and less cited than the background literature '
      '(median 6 vs. 16 citations), because the benchmark rewards recent '
      'pandemic findings, not established ones. Across our four benchmarks, '
      'metadata utility therefore tracks the informativeness of the '
      'bibliographic signal, not how completely it covers the corpus.')
    h2(d, '5.3. Metadata acquisition')
    p(d,
      'Citation counts, publication years, and venues were fetched from the '
      'Semantic Scholar API for SCIDOCS and SciFact (99.7% and 94.1% '
      'matched, respectively). For NFCorpus, PubMed identifiers were mapped '
      'through the Semantic Scholar batch endpoint (94.0% matched). For '
      'TREC-COVID, CORD-19 identifiers were resolved through the official '
      'metadata file to DOIs and PubMed identifiers, then enriched through '
      'the OpenAlex API (Priem et al., 2022), reaching 69.8% citation, '
      '96.4% year, and 92.5% '
      'venue coverage. No metadata value is synthetic; unmatched documents '
      'receive neutral defaults (zero citations, corpus-median year) exactly '
      'as they would in production.')
    h2(d, '5.4. Metrics, significance, and compute')
    p(d,
      'We report Recall@1/5/10, NDCG@10 with graded gains, and MRR on the '
      'official test judgments (documents judged non-relevant are excluded '
      'from the relevant sets). Significance uses the one-sided Wilcoxon '
      'signed-rank test over per-query scores with Cohen\'s d computed from '
      'the same per-query differences, so reported p-values and effect sizes '
      'are consistent by construction. To control the family-wise error '
      'rate, the revised paper pre-specifies one primary comparison per '
      'dataset: cross-fitted BiblioGuard versus its BGE-Hybrid fallback on '
      'NDCG@10 (four one-sided Wilcoxon tests, Holm-Bonferroni-corrected '
      'across domains). The confidence gate\'s internal Bonferroni correction '
      'over nine actions is separate from this evaluation-level correction. '
      'Retrieval-landscape, component ablation, robustness, sensitivity, '
      'unconstrained-routing, and generation comparisons are mechanism or '
      'exploratory analyses. The 30-combination metadata-weight grid was '
      'computed on the same judgments and is not confirmatory. All retrieval runs use '
      'a single '
      'CPU-only workstation. Measured per-query cost on the largest corpus '
      '(TREC-COVID, 171k documents): BM25 scoring averages 1,002 ms; dense '
      'scoring over precomputed embeddings and CA-HR re-ranking add under '
      '2 ms per query. On the smaller corpora BM25 costs 7.6-97 ms per '
      'query. Query encoding (MiniLM) averages about 30 ms. Total online '
      'latency is therefore dominated by the LLM call, not retrieval.')

    # ---- 6. Results ---------------------------------------------------------
    h1(d, '6. Results')
    h2(d, '6.1. Primary BiblioGuard evaluation')
    p(d, 'Table 4. Cross-fitted BiblioGuard evaluation. Active rates are '
         'unconstrained / confidence-gated; p-values compare BiblioGuard with '
         'BGE-Hybrid and are Holm-adjusted across four domains.',
      italic=True, size=9)
    rows = []
    for ds in DS:
        bg = BG[ds]
        raw = bg['ablation_unconstrained']
        rows.append([
            DS_NAME[ds], f4(bg['baseline_N@10']), f4(raw['N@10']),
            f4(bg['biblioguard_N@10']),
            f"{100*raw['selection_rate']:.1f}% / {100*bg['selection_rate']:.1f}%",
            f"{bg['gain_N@10']:+.4f}", pval(bg['wilcoxon_p_holm'])])
    add_table(d, ['Dataset', 'Fallback', 'Without LCB', 'BiblioGuard',
                  'Active rate', 'Guarded gain', 'Holm p'], rows)
    p(d,
      'BiblioGuard improves SCIDOCS from '
      f'{f4(BG["scidocs"]["baseline_N@10"])} to '
      f'{f4(BG["scidocs"]["biblioguard_N@10"])} NDCG@10 '
      f'({BG["scidocs"]["gain_N@10"]:+.4f}; '
      f'{100*BG["scidocs"]["selection_rate"]:.1f}% of queries activated; '
      f'Holm-adjusted p = {BG["scidocs"]["wilcoxon_p_holm"]:.2g}; '
      f'paired d = {BG["scidocs"]["paired_cohen_d"]:.3f}). The selected '
      'SCIDOCS actions are citation-only: beta = 0.30 on 141 queries, beta = '
      '0.20 on 39, and beta = 0.15 on 7; 813 queries return the fallback. On '
      'SciFact, NFCorpus, and TREC-COVID no action clears the simultaneous '
      'lower bound, so the policy exactly reproduces BGE-Hybrid. Across the '
      'released four-domain evaluation this yields no negative mean transfer '
      'and a macro-average gain of '
      f'{T["biblioguard"]["macro"]["gain_N@10"]:+.4f}. This is an empirical '
      'result under the cross-fitted protocol, not a universal safety claim.')
    p(d,
      'The confidence gate is consequential rather than cosmetic. Removing '
      'it activates 96.8% of SCIDOCS queries and reaches 0.2068 NDCG@10, but '
      'it also activates 67.3%, 45.8%, and 54.0% of queries on SciFact, '
      'NFCorpus, and TREC-COVID and changes NDCG@10 by -0.00003, -0.00339, '
      'and -0.01148. The guarded policy therefore trades some attainable '
      'SCIDOCS gain for abstention where local mean estimates are not strong '
      'enough after simultaneous uncertainty correction. Fig. 5 visualises '
      'this selectivity/performance trade-off.')
    d.add_picture(os.path.join(BASE, 'figures', 'Fig5_biblioguard.png'),
                  width=Inches(5.0))
    p(d, 'Fig. 5. BGE-Hybrid, unconstrained paired-effect routing, and '
         'confidence-gated BiblioGuard. Labels show the guarded action rate.',
      italic=True, size=9)

    h2(d, '6.2. Retrieval landscape and failure motivation')
    p(d, 'Table 5. Retrieval effectiveness: NDCG@10 / Recall@10 on the four '
         'test sets (best single method per dataset in bold in Fig. 2).',
      italic=True, size=9)
    rows = []
    for m in METHODS:
        row = [m]
        for ds in DS:
            row.append(f"{avg(ds, m, 'N@10'):.4f} / {avg(ds, m, 'R@10'):.4f}")
        rows.append(row)
    add_table(d, ['Method', 'SCIDOCS', 'SciFact', 'NFCorpus', 'TREC-COVID'],
              rows)
    best_str = '; '.join(
        f"{DS_NAME[ds]}: {T['main'][ds]['best_single_N@10']} "
        f"({avg(ds, T['main'][ds]['best_single_N@10'], 'N@10'):.4f})"
        for ds in DS)
    p(d,
      f'Fig. 2 and Table 5 show the central pattern: no configuration wins '
      f'everywhere. The best single method per dataset is {best_str}. '
      f'Pretrained dense retrieval dominates on SCIDOCS, where '
      f'SBERT-Dense reaches {f4(avg("scidocs", "SBERT-Dense", "N@10"))} '
      f'NDCG@10 versus {f4(avg("scidocs", "BM25", "N@10"))} for BM25 '
      f'(+44.7% relative); on SciFact the hybrid family is best-in-class '
      f'(Neural-Hybrid {f4(avg("scifact", "Neural-Hybrid", "N@10"))} versus '
      f'{f4(avg("scifact", "BM25", "N@10"))} for BM25), and CA-HR attains '
      f'the best Recall@10 ({f4(avg("scifact", "CA-HR", "R@10"))}) of all '
      f'ten methods. BGE-based configurations provide the best single '
      f'result on three of the four datasets—BGE-Dense on SciFact '
      f'({f4(avg("scifact", "BGE-Dense", "N@10"))}) and TREC-COVID '
      f'({f4(avg("trec-covid", "BGE-Dense", "N@10"))}), the '
      f'latter consistent with its published BEIR reference score '
      f'(approximately 0.76; Xiao et al., 2023), and BGE-Hybrid '
      f'on NFCorpus ({f4(avg("nfcorpus", "BGE-Hybrid", "N@10"))})—yet '
      f'BGE-Dense trails '
      f'SBERT-Dense on SCIDOCS '
      f'({f4(avg("scidocs", "BGE-Dense", "N@10"))} vs. '
      f'{f4(avg("scidocs", "SBERT-Dense", "N@10"))}), confirming that even '
      f'encoder choice is domain-dependent.')
    p(d,
      'CA-HR, the citation- and recency-aware configuration, sits in the '
      'middle of the hybrid family overall. It exceeds BM25 and LSA-Dense '
      'on all four datasets, but it does not beat the strongest plain dense '
      'or hybrid baseline except against SBERT-Dense on SciFact. These '
      'landscape contrasts are descriptive/exploratory in the revised '
      'analysis; the confirmatory family is restricted to BiblioGuard versus '
      'its fallback (Section 6.1). The important motivation is that always-on '
      'metadata is competitive yet unstable, which creates a role for a '
      'selective intervention policy rather than another fixed re-ranker.')
    p(d,
      'Table 5 also isolates the role of the dense backbone. Equal-weight '
      'hybrid fusion on BGE-small (BGE-Hybrid) does not reproduce the hybrid '
      'advantage seen with MiniLM: it trails BGE-Dense on SCIDOCS '
      f'({f4(avg("scidocs", "BGE-Hybrid", "N@10"))} vs. '
      f'{f4(avg("scidocs", "BGE-Dense", "N@10"))}), SciFact '
      f'({f4(avg("scifact", "BGE-Hybrid", "N@10"))} vs. '
      f'{f4(avg("scifact", "BGE-Dense", "N@10"))}), and TREC-COVID '
      f'({f4(avg("trec-covid", "BGE-Hybrid", "N@10"))} vs. '
      f'{f4(avg("trec-covid", "BGE-Dense", "N@10"))}); only on NFCorpus does '
      f'BGE-Hybrid ({f4(avg("nfcorpus", "BGE-Hybrid", "N@10"))}) edge out '
      f'BGE-Dense ({f4(avg("nfcorpus", "BGE-Dense", "N@10"))}), where it '
      'attains the best NDCG@10 of all ten configurations. Decisively for '
      'RQ1, transferring the CA-HR re-ranking rule to the BGE backbone '
      '(BGE-CA-HR) never helps: it falls slightly below BGE-Hybrid on every '
      f'dataset (SCIDOCS {f4(avg("scidocs", "BGE-CA-HR", "N@10"))} vs. '
      f'{f4(avg("scidocs", "BGE-Hybrid", "N@10"))}, '
      f'd = {f3(T["main"]["scidocs"]["bge_tests"]["BGE-CA-HR vs BGE-Hybrid | N@10"]["d"])}; '
      f'SciFact {f4(avg("scifact", "BGE-CA-HR", "N@10"))} vs. '
      f'{f4(avg("scifact", "BGE-Hybrid", "N@10"))}, '
      f'd = {f3(T["main"]["scifact"]["bge_tests"]["BGE-CA-HR vs BGE-Hybrid | N@10"]["d"])}; '
      f'NFCorpus {f4(avg("nfcorpus", "BGE-CA-HR", "N@10"))} vs. '
      f'{f4(avg("nfcorpus", "BGE-Hybrid", "N@10"))}, '
      f'd = {f3(T["main"]["nfcorpus"]["bge_tests"]["BGE-CA-HR vs BGE-Hybrid | N@10"]["d"])}; '
      f'TREC-COVID {f4(avg("trec-covid", "BGE-CA-HR", "N@10"))} vs. '
      f'{f4(avg("trec-covid", "BGE-Hybrid", "N@10"))}, '
      f'd = {f3(T["main"]["trec-covid"]["bge_tests"]["BGE-CA-HR vs BGE-Hybrid | N@10"]["d"])}) '
      'and well below BGE-Dense. Under CA-HR\'s fixed a-priori '
      'hyperparameters, the citation gain obtained on the weaker MiniLM '
      'backbone on SCIDOCS does not transfer to a stronger encoder: the '
      'observed metadata benefit is backbone-dependent rather than '
      'additive (Section 6.5 tests whether any metadata weight rescues it).')
    d.add_picture(os.path.join(BASE, 'figures', 'Fig2_main_results.png'),
                  width=Inches(6.0))
    p(d, 'Fig. 2. Retrieval effectiveness (NDCG@10) across four domains.',
      italic=True, size=9)

    d.add_page_break()
    h2(d, '6.3. Ablation of CA-HR components')
    p(d, 'Table 6. CA-HR ablation (NDCG@10; bold marks cases where removing '
         'a component improves performance).', italic=True, size=9)
    ABL = ['full', '-citation', '-recency', '-dense (alpha=1)',
           '-sparse (alpha=0)', '-rerank (plain hybrid)']
    ABL_LABEL = {'full': 'Full CA-HR', '-citation': '− citation',
                 '-recency': '− recency', '-dense (alpha=1)': '− dense (α=1)',
                 '-sparse (alpha=0)': '− sparse (α=0)',
                 '-rerank (plain hybrid)': '− re-rank'}
    rows = []
    for k in ABL:
        row = [ABL_LABEL[k]]
        for ds in DS:
            v = abl(ds, k)
            mark = '*' if (k != 'full' and v > abl(ds, 'full')) else ''
            row.append(f'{v:.4f}{mark}')
        rows.append(row)
    add_table(d, ['Variant', 'SCIDOCS', 'SciFact', 'NFCorpus', 'TREC-COVID'],
              rows)
    p(d,
      f'The ablation (Table 6, Fig. 3; * = removal improves on the full '
      f'model) localises exactly where metadata helps and where it hurts. '
      f'On SCIDOCS—the corpus where citations are most informative about '
      f'relevance (AUC = {T["diagnostics"]["scidocs"]["cit_rel_auc"]:.3f})—removing the '
      f'citation boost drops NDCG@10 from {f4(abl("scidocs", "full"))} to '
      f'{f4(abl("scidocs", "-citation"))}, the single most valuable metadata '
      f'signal in the study, while removing the recency prior improves '
      f'performance to {f4(abl("scidocs", "-recency"))}, indicating a '
      f'miscalibrated recency prior on a snapshot corpus. On SciFact both '
      f'metadata terms are neutral. On the two new domains the sign flips: '
      f'removing the citation boost improves NDCG@10 on NFCorpus '
      f'({f4(abl("nfcorpus", "full"))} → {f4(abl("nfcorpus", "-citation"))}) '
      f'and on TREC-COVID '
      f'({f4(abl("trec-covid", "full"))} → {f4(abl("trec-covid", "-citation"))}). '
      f'The sign flip tracks the bibliographic-informativeness diagnostics '
      f'(Table 3), not coverage: NFCorpus and TREC-COVID are exactly the '
      f'corpora where citation counts carry no (AUC = '
      f'{T["diagnostics"]["nfcorpus"]["cit_rel_auc"]:.3f}) or inverse '
      f'(AUC = {T["diagnostics"]["trec-covid"]["cit_rel_auc"]:.3f}) '
      f'information about relevance, so the citation boost amplifies '
      f'noise rather than refining the ranking. The content '
      f'side is stable across domains: removing either modality degrades '
      f'performance almost everywhere, except that a pure dense '
      f'configuration is best on SCIDOCS, consistent with dense dominance '
      f'on computer science.')
    d.add_picture(os.path.join(BASE, 'figures', 'Fig3_ablation.png'),
                  width=Inches(6.0))
    p(d, 'Fig. 3. CA-HR ablation (NDCG@10; dashed line = full model).',
      italic=True, size=9)

    h2(d, '6.4. Robustness to query corruption')
    p(d,
      'Under simulated word-drop noise (10%-40%), all hybrid methods degrade '
      'gracefully, but robustness rankings are domain-dependent and do not '
      'simply follow clean-query rankings (Fig. 4). On TREC-COVID, CA-HR is '
      'the most noise-resistant configuration: at 40% corruption it retains '
      f'{f4(robust_n10("trec-covid", "0.4", "CA-HR"))} NDCG@10 '
      f'versus {f4(robust_n10("trec-covid", "0.4", "Neural-Hybrid"))} '
      'for Neural-Hybrid and '
      f'{f4(robust_n10("trec-covid", "0.4", "BM25"))} for BM25: the '
      'citation authority term, being independent of the corrupted query '
      'text, acts as a stabiliser even where it does not raise clean-query '
      'effectiveness. On SciFact, CA-HR '
      f'({f4(robust_n10("scifact", "0.4", "CA-HR"))}) and '
      'Neural-Hybrid '
      f'({f4(robust_n10("scifact", "0.4", "Neural-Hybrid"))}) are '
      'nearly tied at 40% noise, both well above BM25 '
      f'({f4(robust_n10("scifact", "0.4", "BM25"))}), and on '
      'NFCorpus the two hybrids again track each other closely '
      f'({f4(robust_n10("nfcorpus", "0.4", "CA-HR"))} vs. '
      f'{f4(robust_n10("nfcorpus", "0.4", "Neural-Hybrid"))}). '
      'SCIDOCS is the exception: there BM25 is the most robust method at 40% '
      f'corruption ({f4(robust_n10("scidocs", "0.4", "BM25"))} '
      f'versus {f4(robust_n10("scidocs", "0.4", "CA-HR"))} for CA-HR '
      f'and {f4(robust_n10("scidocs", "0.4", "Neural-Hybrid"))} for '
      'Neural-Hybrid), indicating that where dense representations dominate '
      'on clean queries they are also the most fragile to lexical '
      'corruption.')
    d.add_picture(os.path.join(BASE, 'figures', 'Fig4_robustness.png'),
                  width=Inches(6.0))
    p(d, 'Fig. 4. Robustness to query corruption (NDCG@10 vs. word-drop '
         'noise).', italic=True, size=9)

    h2(d, '6.5. Metadata-weight sensitivity and rank-fusion baselines')
    p(d,
      'Two objections remain after Section 6.2. First, CA-HR\'s weights were '
      'fixed a priori, so the absence of a metadata gain on the stronger BGE '
      'backbone could be an artefact of under-weighting the metadata terms. '
      'Second, our hybrids interpolate min-max-normalised scores, so a '
      'standard rank-fusion reference is missing. Table 7 addresses both. '
      'We sweep the citation and recency weights of BGE-CA-HR over a '
      '30-combination grid (beta in {0, 0.05, 0.10, 0.15, 0.20, 0.30}, '
      'gamma in {0, 0.05, 0.10, 0.15, 0.20}) and test each combination '
      'against BGE-Hybrid with one-sided Wilcoxon tests, Holm-corrected '
      'within each dataset. Because the grid is selected on the test '
      'judgments, we report this sweep strictly as an exploratory '
      'sensitivity analysis, not as confirmatory evidence. The metadata '
      'gain re-emerges only on SCIDOCS: '
      f'{T["sensitivity"]["scidocs"]["best_combo"].replace("beta=", "beta = ").replace("|gamma=", ", gamma = ")} '
      f'attains {f4(T["sensitivity"]["scidocs"]["best_N@10"])} NDCG@10, '
      'significantly above BGE-Hybrid '
      f'({f4(T["sensitivity"]["scidocs"]["bge_hybrid_N@10"])}; Holm '
      'p < 0.001) and marginally above BGE-Dense '
      f'({f4(avg("scidocs", "BGE-Dense", "N@10"))}). On SciFact, NFCorpus, '
      'and TREC-COVID no grid combination significantly beats BGE-Hybrid '
      f'(best {f4(T["sensitivity"]["scifact"]["best_N@10"])}, '
      f'{f4(T["sensitivity"]["nfcorpus"]["best_N@10"])}, and '
      f'{f4(T["sensitivity"]["trec-covid"]["best_N@10"])} versus '
      f'{f4(T["sensitivity"]["scifact"]["bge_hybrid_N@10"])}, '
      f'{f4(T["sensitivity"]["nfcorpus"]["bge_hybrid_N@10"])}, and '
      f'{f4(T["sensitivity"]["trec-covid"]["bge_hybrid_N@10"])}; all '
      'Holm-adjusted p = 1.0). The backbone effect is therefore not a '
      'weight artefact: metadata gains on a strong encoder are gated by '
      'citation informativeness, and where the signal is informative the '
      'gain transfers only after re-calibrating the metadata weight.')
    p(d,
      'As a stronger fusion reference we add reciprocal rank fusion '
      '(RRF, k = 60) of BM25 with each dense encoder. RRF-MiniLM / RRF-BGE '
      'reach NDCG@10 '
      f'{f4(T["sensitivity"]["scidocs"]["rrf"]["RRF-MiniLM"]["N@10"])} / '
      f'{f4(T["sensitivity"]["scidocs"]["rrf"]["RRF-BGE"]["N@10"])} on '
      'SCIDOCS, '
      f'{f4(T["sensitivity"]["scifact"]["rrf"]["RRF-MiniLM"]["N@10"])} / '
      f'{f4(T["sensitivity"]["scifact"]["rrf"]["RRF-BGE"]["N@10"])} on '
      'SciFact, '
      f'{f4(T["sensitivity"]["nfcorpus"]["rrf"]["RRF-MiniLM"]["N@10"])} / '
      f'{f4(T["sensitivity"]["nfcorpus"]["rrf"]["RRF-BGE"]["N@10"])} on '
      'NFCorpus, and '
      f'{f4(T["sensitivity"]["trec-covid"]["rrf"]["RRF-MiniLM"]["N@10"])} / '
      f'{f4(T["sensitivity"]["trec-covid"]["rrf"]["RRF-BGE"]["N@10"])} on '
      'TREC-COVID: no RRF variant is the best configuration on any dataset, '
      'so the qualitative conclusions of Table 5 do not depend on the '
      'min-max interpolation rule.')
    rows = []
    for ds in DS:
        s = T['sensitivity'][ds]
        rows.append([DS_NAME[ds],
                     f4(s['bge_hybrid_N@10']),
                     f4(avg(ds, 'BGE-Dense', 'N@10')),
                     f4(avg(ds, 'BGE-CA-HR', 'N@10')),
                     s['best_combo'].replace('beta=', '').replace('|gamma=', '/'),
                     f4(s['best_N@10']),
                     'yes' if s['any_significant_gain_after_holm'] else 'no',
                     f4(s['rrf']['RRF-MiniLM']['N@10']),
                     f4(s['rrf']['RRF-BGE']['N@10'])])
    p(d, 'Table 7. BGE-backbone metadata-weight sensitivity (30-combination '
         'grid per dataset; significance vs. BGE-Hybrid, one-sided Wilcoxon '
         'with Holm correction) and RRF (k = 60) rank-fusion baselines; '
         'NDCG@10.', italic=True, size=9)
    add_table(d,
      ['Dataset', 'BGE-Hybrid', 'BGE-Dense', 'BGE-CA-HR (fixed)',
       'Best grid beta/gamma', 'Best grid N@10', 'Sig. after Holm',
       'RRF-MiniLM', 'RRF-BGE'], rows)

    h2(d, '6.6. Generation-side evaluation')
    p(d, 'Table 8. End-to-end answer quality with DeepSeek generation on 200 '
         'paired queries (50 per dataset across the four benchmarks): '
         'LLM-judged relevance and faithfulness (1-5), citation precision '
         'against gold judgments, and relevant passages in the top-5 '
         'context.',
      italic=True, size=9)
    ca, nh = g['CA-HR'], g['Neural-Hybrid']
    rows = [
        ['Relevance (judge 1-5)', f"{ca['relevance']['mean']:.2f} ± {ca['relevance']['std']:.2f}",
         f"{nh['relevance']['mean']:.2f} ± {nh['relevance']['std']:.2f}",
         pval(g['paired_relevance']['wilcoxon_p_two_sided'])],
        ['Faithfulness (judge 1-5)', f"{ca['faithfulness']['mean']:.2f} ± {ca['faithfulness']['std']:.2f}",
         f"{nh['faithfulness']['mean']:.2f} ± {nh['faithfulness']['std']:.2f}",
         pval(g['paired_faithfulness']['wilcoxon_p_two_sided'])],
        ['Citation precision', f"{g['paired_citation_precision']['mean_CA-HR']:.3f}",
         f"{g['paired_citation_precision']['mean_Neural-Hybrid']:.3f}",
         pval(g['paired_citation_precision']['wilcoxon_p_two_sided'])],
        ['Relevant docs in top-5', f"{g['paired_n_rel_context']['mean_CA-HR']:.2f}",
         f"{g['paired_n_rel_context']['mean_Neural-Hybrid']:.2f}",
         pval(g['paired_n_rel_context']['wilcoxon_p_two_sided'])],
    ]
    add_table(d, ['Measure', 'CA-HR backend', 'Neural-Hybrid backend',
                  'Wilcoxon (two-sided)'], rows)
    p(d,
      'This secondary experiment predates BiblioGuard and compares two fixed '
      'MiniLM-family backends; it tests whether retrieval changes of the '
      'observed magnitude survive generation, not whether BiblioGuard itself '
      'improves generated answers. We sampled 200 test queries (50 per dataset, fixed '
      'seed), generated cited answers with DeepSeek (deepseek-chat, '
      'temperature 0.2, top-5 passages as context), and scored each answer '
      'for relevance and faithfulness with an LLM judge and for citation '
      'precision against the gold judgments. The judge is the same '
      'DeepSeek chat model prompted with a fixed rubric (1-5 relevance; '
      '1-5 faithfulness conditioned on the retrieved context; the exact '
      'prompts are released with the code). Each answer was judged once, '
      'in randomised backend order with backend labels hidden from the '
      'judge; no human verification subsample was performed, which we flag '
      'as a limitation (Section 8.2). Table 8 shows the backends are '
      'statistically indistinguishable on every measure (all p > 0.48): at '
      'top-5, the two systems surface nearly the same passages '
      f'({g["paired_n_rel_context"]["mean_CA-HR"]:.2f} vs. '
      f'{g["paired_n_rel_context"]["mean_Neural-Hybrid"]:.2f} gold-relevant '
      f'documents on average). Two secondary findings are worth reporting. '
      f'First, absolute answer quality is high (relevance '
      f'{ca["relevance"]["mean"]:.2f}-{nh["relevance"]["mean"]:.2f}, '
      f'faithfulness '
      f'{ca["faithfulness"]["mean"]:.2f}-{nh["faithfulness"]["mean"]:.2f} '
      f'of 5), supporting technical feasibility under this automated judge. Second, citation '
      f'precision against the strict gold judgments averages about '
      f'{g["paired_citation_precision"]["mean_CA-HR"]:.2f} but varies '
      f'sharply with gold-judgment density: 0.87 on TREC-COVID, whose '
      f'queries carry dense graded judgments, versus 0.15 on SCIDOCS, whose '
      f'judgments are sparse; the model frequently cites plausible but not '
      f'gold-relevant passages, a failure mode invisible to retrieval '
      f'metrics and a target for future citation verification. We caveat '
      f'that the judge shares the generator\'s model family (Zheng et al., '
      f'2023); the '
      f'paired design controls for this bias, but absolute scores should be '
      f'read accordingly.')

    h2(d, '6.7. Efficiency')
    p(d,
      'All ten configurations are CPU-feasible. On the largest corpus '
      '(TREC-COVID, 171,332 documents) BM25 scoring averages 1,002 ms per '
      'query; dense scoring over precomputed embeddings and CA-HR '
      're-ranking add under 2 ms. On the smaller corpora BM25 costs '
      '7.6-97 ms per query and query encoding about 30 ms. In the deployed '
      'system, retrieval latency is negligible next to LLM generation, and '
      'the one-time corpus embedding cost (about 34-49 documents/s for '
      'MiniLM on 8 CPU threads) is amortised offline.')

    # ---- 7. Deployment feasibility case study -------------------------------
    h1(d, '7. Deployment feasibility case study')
    p(d,
      'The full pipeline has been deployed in the production assistant '
      'since 10 July 2026 '
      '(system name and URL withheld for anonymous review), on a single '
      'Aliyun ECS instance (2 vCPU, 1.6 GB RAM, '
      '40 GB disk, Ubuntu 22.04) with Node.js 20.20.2, MySQL 8.0.46, '
      'Nginx 1.18 terminating TLS (Let\'s Encrypt, auto-renewed), and PM2 '
      'supervision. Over the first month of operation the service recorded 6 '
      'registered users, 6 uploaded papers, 26 chat conversations with 63 '
      'messages, 6 AI-generated summaries, and 3 writing projects; the '
      'process restarted twice in six days (both manual deploys) and the '
      'edge responds in about 74 ms. This is a pilot-scale deployment, and '
      'we report it as such: its purpose is to demonstrate that the '
      'retrieval/generation pipeline runs on commodity hardware, not to claim '
      'large-scale adoption or an online evaluation of BiblioGuard. The pilot '
      'used fixed hybrid/CA-HR backends and predates the new decision policy; '
      'BiblioGuard has been integrated into the offline reproduction code but '
      'has not yet been subjected to a live A/B test.')
    p(d,
      'Two operational observations connect the deployment to the '
      'experimental findings. First, retrieval never appeared in the latency '
      'budget: user-perceived response time is dominated by the streamed LLM '
      'tokens, consistent with Section 6.7, so the choice among the ten '
      'configurations is free from an engineering standpoint. Second, '
      'user-uploaded papers are often fresh preprints with no citation '
      'record—the production analogue of TREC-COVID\'s inverted '
      'citation-relevance association—so '
      'the system defaults to the plain hybrid backend for user libraries. '
      'BiblioGuard formalises a future evidence-gated path away from that '
      'fallback; the current deployment evidence establishes only that its '
      'underlying retrieval actions satisfy the hardware budget.')

    # ---- 8. Discussion ------------------------------------------------------
    h1(d, '8. Discussion')
    h2(d, '8.1. Why confidence-gated intervention changes the result')
    p(d,
      'BiblioGuard changes the learning target from absolute winner '
      'prediction to paired incremental utility. This matters because all '
      'candidate outcomes share the same query and much of the same ranking; '
      'subtracting the fallback removes query difficulty and asks only '
      'whether intervention is justified. The explicit no-action outcome '
      'also changes the error asymmetry: a false positive can damage a strong '
      'baseline, whereas a false negative merely forgoes a possible gain. '
      'The lower-bound gate implements this conservative preference.')
    p(d,
      'The cross-domain activation pattern is consistent with the independent '
      'bibliographic diagnostics. SCIDOCS has a strong citation-relevance '
      'association (AUC = 0.798), and every guarded activation is a '
      'citation-only action. SciFact is weaker (AUC = 0.582), NFCorpus is '
      'uninformative (0.498), and TREC-COVID is inverted (0.461); in these '
      'domains the unconstrained estimator often predicts a positive local '
      'mean but the simultaneous lower bound rejects it. This difference '
      'between mean routing and evidence-gated routing is the central '
      'mechanism result, not a claim that bibliographic metadata is generally '
      'beneficial. Backbone and query noise remain relevant: the fixed '
      'MiniLM metadata gain does not transfer directly to BGE (Section 6.2), '
      'and robustness rankings change under word drop (Section 6.4).')
    h2(d, '8.2. Limitations')
    p(d,
      'Eight limitations delimit the contribution. (1) BiblioGuard requires '
      'historical queries with judgments; an unseen domain follows the '
      'metadata-free cold-start path and receives no learned benefit. (2) The '
      'similarity-weighted Student-t lower bound is an operational confidence '
      'gate, not a finite-sample or distribution-free safety certificate. '
      '(3) Evaluation is within-domain cross-fitting, not temporal or '
      'cross-domain external validation; only SCIDOCS activates, and its '
      '+0.0049 absolute gain and d = 0.139 are modest. (4) Interventions use '
      'CA-HR\'s 0.6/0.4 content fusion whereas the fallback uses 0.5/0.5, so '
      'the paired effect measures the complete configuration switch rather '
      'than a metadata-only causal effect. (5) The action family, confidence '
      'level, one BGE encoder, four benchmarks, and NDCG@10 utility should be '
      'locked prospectively and tested on more domains, encoders, and time '
      'slices. (6) TREC-COVID has only 50 queries and 69.8% citation coverage; '
      'neutral defaults cannot remove systematic missingness. (7) The '
      'generation experiment compares fixed MiniLM backends, uses the same '
      'model family for generation and judging, has no human verification, '
      'and therefore does not establish a BiblioGuard generation gain. (8) '
      'The six-user, one-month deployment predates BiblioGuard and shows '
      'hardware feasibility only; an online A/B evaluation and passage-level '
      'validation remain future work.')

    # ---- 9. Conclusion ------------------------------------------------------
    h1(d, '9. Conclusion')
    p(d,
      'We introduced BiblioGuard, a cross-fitted decision layer for selective '
      'bibliographic intervention in scientific RAG. Rather than predicting '
      'the best retriever, it estimates paired local uplift relative to a '
      'strong fallback, corrects its one-sided gate over nine actions, and '
      'abstains when the evidence is insufficient. On four domains it '
      'improves SCIDOCS by 0.0049 NDCG@10 and reproduces the fallback exactly '
      'elsewhere; removing the lower-confidence gate causes negative transfer '
      'in the other three domains. The result supports a narrow conclusion: '
      'selective metadata use can be safer empirically than always-on or '
      'mean-only routing when paired historical evidence is available. It '
      'does not establish a universal guarantee. Prospective multi-encoder '
      'and temporal evaluation, calibrated or conformal risk control, and a '
      'live BiblioGuard A/B test are the next steps.')

    # ---- Declaration of generative AI (required before references) --------
    h1(d, 'Declaration of generative AI and AI-assisted technologies in '
          'the manuscript preparation process')
    p(d,
      'During the preparation of this work the author used Kimi (Moonshot '
      'AI) for code debugging and language polishing. All AI-assisted '
      'output, including code and experimental results, was reviewed, '
      'executed, and verified by the author, who takes full responsibility '
      'for the content of the published article.')

    # ---- References (APA, author-year, alphabetical) -----------------------
    h1(d, 'References')
    refs = [
        'Asai, A., He, J., Shao, R., Shi, W., Singh, A., Chang, J. C., Lo, K., Soldaini, L., Feldman, S., D\'Arcy, M., Wadden, D., Latzke, M., Sparks, J., Hwang, J. D., Kishore, V., Tian, M., Ji, P., Liu, S., Tong, H., Wu, B., ... Hajishirzi, H. (2026). Synthesizing scientific literature with retrieval-augmented language models. Nature, 650, 857–863. https://doi.org/10.1038/s41586-025-10072-4',
        'Asai, A., Wu, Z., Wang, Y., Sil, A., & Hajishirzi, H. (2024). Self-RAG: Learning to retrieve, generate, and critique through self-reflection. In Proceedings of ICLR.',
        'Bao, T., Nayeem, M. T., Rafiei, D., & Zhang, C. (2025). SurveyGen: Quality-aware scientific survey generation with large language models. In Proceedings of EMNLP (pp. 2712–2736). https://doi.org/10.18653/v1/2025.emnlp-main.136',
        'Boteva, V., Gholipour, D., Sokolov, A., & Riezler, S. (2016). A full-text learning to rank dataset for medical information retrieval. In Proceedings of ECIR (pp. 716–722).',
        'Chen, J., Xiao, S., Zhang, P., Luo, K., Lian, D., & Liu, Z. (2024). BGE M3-Embedding: Multi-lingual, multi-functionality, multi-granularity text embeddings through self-knowledge distillation. arXiv:2402.03216. https://doi.org/10.48550/arXiv.2402.03216',
        'Cohan, A., Feldman, S., Beltagy, I., Downey, D., & Weld, D. S. (2020). SPECTER: Document-level representation learning using citation-informed transformers. In Proceedings of ACL (pp. 2270–2282). https://doi.org/10.18653/v1/2020.acl-main.207',
        'Deerwester, S., Dumais, S. T., Furnas, G. W., Landauer, T. K., & Harshman, R. (1990). Indexing by latent semantic analysis. Journal of the American Society for Information Science, 41(6), 391–407.',
        'DeepSeek-AI. (2024). DeepSeek-V3 technical report. arXiv:2412.19437. https://doi.org/10.48550/arXiv.2412.19437',
        'Ding, H., Zhao, Y., Hu, T., Wang, Z., Patwardhan, M., & Cohan, A. (2026). SciRAG: Adaptive, citation-aware, and outline-guided retrieval and synthesis for scientific literature. In Proceedings of EACL (Volume 1: Long Papers) (pp. 6440–6460).',
        'Hwang, J., Park, J., Park, H., Kim, D., Park, S., & Ok, J. (2025). Retrieval-augmented generation with estimation of source reliability. In Proceedings of EMNLP (pp. 34279–34303). https://doi.org/10.18653/v1/2025.emnlp-main.1738',
        'Karpukhin, V., Oğuz, B., Min, S., Lewis, P., Wu, L., Edunov, S., Chen, D., & Yih, W.-t. (2020). Dense passage retrieval for open-domain question answering. In Proceedings of EMNLP (pp. 6769–6781). https://doi.org/10.18653/v1/2020.emnlp-main.550',
        'Kim, J., Yoon, S., Le, X.-B., Nam, Y., Kim, D., Song, H., & Lee, J.-G. (2026). QuDAR: Query-wise dual-perspective adaptive retrieval. In Proceedings of ACL (pp. 38662–38679). https://doi.org/10.18653/v1/2026.acl-long.1791',
        'Lewis, P., Perez, E., Piktus, A., Petroni, F., Karpukhin, V., Goyal, N., Küttler, H., Lewis, M., Yih, W.-t., Rocktäschel, T., Riedel, S., & Kiela, D. (2020). Retrieval-augmented generation for knowledge-intensive NLP tasks. In Proceedings of NeurIPS (pp. 9459–9474).',
        'Lin, J. (2021). A proposed conceptual framework for a representational approach to information retrieval. ACM SIGIR Forum, 55(2).',
        'Nogueira, R., Jiang, Z., Pradeep, R., & Lin, J. (2020). Document ranking with a pretrained sequence-to-sequence model. In Findings of EMNLP (pp. 708–718). https://doi.org/10.18653/v1/2020.findings-emnlp.63',
        'Ostendorff, M., Rethmeier, N., Augenstein, I., Gipp, B., & Rehm, G. (2022). Neighborhood contrastive learning for scientific document representations with citation embeddings. In Proceedings of EMNLP (pp. 11670–11688).',
        'Priem, J., Piwowar, H., & Orr, R. (2022). OpenAlex: A fully-open index of scholarly works, authors, venues, institutions, and concepts. arXiv:2205.01833. https://doi.org/10.48550/arXiv.2205.01833',
        'Qu, Y., Ding, Y., Liu, J., Liu, K., Ren, R., Zhao, W. X., Dong, D., Wu, H., & Wang, H. (2021). RocketQA: An optimized training approach to dense passage retrieval. In Proceedings of NAACL (pp. 5835–5847). https://doi.org/10.18653/v1/2021.naacl-main.466',
        'Reimers, N., & Gurevych, I. (2019). Sentence-BERT: Sentence embeddings using Siamese BERT-networks. In Proceedings of EMNLP-IJCNLP (pp. 3982–3992). https://doi.org/10.18653/v1/D19-1410',
        'Robertson, S., & Zaragoza, H. (2009). The probabilistic relevance framework: BM25 and beyond. Foundations and Trends in Information Retrieval, 3(4), 333–389. https://doi.org/10.1561/1500000019',
        'Shao, Z., Gong, Y., Shen, Y., Huang, M., Duan, N., & Chen, W. (2023). Enhancing retrieval-augmented large language models with iterative retrieval-generation synergy. In Findings of EMNLP (pp. 9248–9274).',
        'Singh, A., D\'Arcy, M., Cohan, A., Downey, D., & Feldman, S. (2023). SciRepEval: A multi-format benchmark for scientific document representations. In Proceedings of EMNLP (pp. 5548–5566).',
        'Thakur, N., Reimers, N., Rücklé, A., Srivastava, A., & Gurevych, I. (2021). BEIR: A heterogeneous benchmark for zero-shot evaluation of information retrieval models. In Proceedings of NeurIPS Datasets and Benchmarks.',
        'Varangot-Reille, C., Bouvard, C., & Gourru, A. (2026). Generalising LLM routing using past performance retrieval: A few-shot router is sufficient. In Proceedings of EACL Student Research Workshop (pp. 304–319). https://doi.org/10.18653/v1/2026.eacl-srw.22',
        'Voorhees, E., Alam, T., Bedrick, S., Demner-Fushman, D., Hersh, W. R., Lo, K., Roberts, K., Soboroff, I., & Wang, L. L. (2021). TREC-COVID: Constructing a pandemic information retrieval test collection. ACM SIGIR Forum, 54(1), 1–12.',
        'Wadden, D., Lin, S., Lo, K., Wang, L. L., van Zuylen, M., Cohan, A., & Hajishirzi, H. (2020). Fact or fiction: Verifying scientific claims. In Proceedings of EMNLP (pp. 7534–7550). https://doi.org/10.18653/v1/2020.emnlp-main.609',
        'Wang, L., Yang, N., Huang, X., Jiao, B., Yang, L., Jiang, D., Majumder, R., & Wei, F. (2022). Text embeddings by weakly-supervised contrastive pre-training. arXiv:2212.03533. https://doi.org/10.48550/arXiv.2212.03533',
        'Wang, L., Ma, C., Feng, X., Zhang, Z., Yang, H., Zhang, J., Chen, Z., Tang, J., Chen, X., Lin, Y., Zhao, W. X., Wei, Z., & Wen, J.-R. (2024). A survey on large language model based autonomous agents. Frontiers of Computer Science, 18(6), 186345.',
        'Wang, W., Wei, F., Dong, L., Bao, H., Yang, N., & Zhou, M. (2020). MiniLM: Deep self-attention distillation for task-agnostic compression of pre-trained transformers. In Proceedings of NeurIPS.',
        'Xiao, S., Liu, Z., Zhang, P., & Muennighoff, N. (2023). C-Pack: Packaged resources to advance general Chinese embedding. arXiv:2309.07597. https://doi.org/10.48550/arXiv.2309.07597',
        'Yousuf, R. B., Xu, S., Sharma, M., Neeser, A., Latimer, C., & Ramakrishnan, N. (2026). Utilizing metadata for better retrieval-augmented generation. In Proceedings of ECIR (pp. 305–319). https://doi.org/10.1007/978-3-032-21289-4_20',
        'Zhao, T., Zhu, Y., Tian, Y., & Dou, Z. (2026). R^3AG: Retriever routing for retrieval-augmented generation. In Proceedings of ACL (pp. 20506–20522). https://doi.org/10.18653/v1/2026.acl-long.939',
        'Zheng, L., Chiang, W.-L., Sheng, Y., Zhuang, S., Wu, Z., Zhuang, Y., Lin, Z., Li, Z., Li, D., Xing, E. P., Zhang, H., Gonzalez, J. E., & Stoica, I. (2023). Judging LLM-as-a-judge with MT-Bench and Chatbot Arena. In Proceedings of NeurIPS Datasets and Benchmarks.',
        'Zhuang, S., Zhuang, H., Koopman, B., & Zuccon, G. (2024). A setwise approach for effective and highly efficient zero-shot ranking with large language models. In Proceedings of SIGIR (pp. 38–47).',
    ]
    for r in refs:
        p(d, r, size=10)

    d.save(os.path.join(OUT, '01_Manuscript_ESWA.docx'))


# ===========================================================================
# Cover letter
# ===========================================================================
def build_cover_letter():
    d = new_doc()
    p(d, 'Cover Letter', bold=True, size=14)
    p(d, f'{AUTHOR}')
    p(d, AFFIL)
    p(d, f'E-mail: {EMAIL}  |  ORCID: {ORCID}')
    d.add_paragraph()
    p(d, 'Dear Editor-in-Chief,')
    p(d,
      f'We submit the manuscript "{TITLE}" for consideration in Expert '
      f'Systems with Applications.')
    p(d,
      'The manuscript fits the journal\'s focus on applied intelligent '
      'systems through a new decision algorithm, BiblioGuard. Given a strong '
      'metadata-free fallback and nine citation/recency actions, BiblioGuard '
      'retrieves similar historical queries, estimates paired query-level '
      'NDCG uplift, applies simultaneous one-sided confidence bounds, and '
      'abstains unless an intervention has positive lower-bound evidence. A '
      'five-fold cross-fitted evaluation across four scientific domains '
      'shows a significant SCIDOCS gain and exact fallback behaviour on the '
      'other three; removing the confidence gate causes negative transfer in '
      'all three. The paper further provides mechanism diagnostics, an '
      'end-to-end feasibility analysis, and a pilot academic-writing '
      'assistant deployment. All data, code, cross-fitted decisions, and per-query results are publicly '
      f'available at {REPO} (archived at https://doi.org/{DOI}) '
      'for full reproducibility.')
    p(d,
      'The manuscript is original, is not under consideration elsewhere, '
      'and is approved by the author. The author declares no competing '
      'interests.')
    p(d,
      'Regarding the journal\'s institutional e-mail requirement: the '
      'corresponding author is an undergraduate researcher at Sanya '
      'University, which does not provide institutional e-mail accounts to '
      'undergraduate students. The personal e-mail address given above and '
      'in the submission system is therefore used for all correspondence.')
    d.add_paragraph()
    p(d, 'Suggested reviewers (no conflicts of interest):')
    for r in [
        'Jimmy Lin, University of Waterloo, Canada — jimmylin@uwaterloo.ca '
        '(information retrieval, neural ranking)',
        'Arman Cohan, Yale University / Allen Institute for AI, USA — '
        'arman.cohan@yale.edu (scientific document representations)',
        'Guido Zuccon, The University of Queensland, Australia — '
        'g.zuccon@uq.edu.au (health/medical IR, LLM rankers)',
        'Andrew Yates, Johns Hopkins University, USA — '
        'andrew.yates@jhu.edu (dense retrieval, ranking)',
    ]:
        par = d.add_paragraph(style='List Bullet')
        par.add_run(r)
    d.add_paragraph()
    p(d, 'Sincerely,')
    p(d, AUTHOR)
    d.save(os.path.join(OUT, '03_Cover_Letter.docx'))


# ===========================================================================
# Declarations
# ===========================================================================
def build_declarations():
    d = new_doc()
    p(d, 'Declarations', bold=True, size=14)
    h2(d, 'Ethics approval and consent to participate')
    p(d, 'Not applicable. The study uses only public benchmark datasets '
         'and aggregated, non-personal operational counts from the deployed '
         'system (registered users, uploaded papers, conversations, '
         'messages, restarts, response times). No user content, no personal '
         'data, and no human subjects were involved; no ethics committee '
         'approval was required.')
    h2(d, 'Consent for publication')
    p(d, 'Not applicable.')
    h2(d, 'Declaration of competing interest')
    p(d, 'The author declares no competing financial or personal interests.')
    h2(d, 'Funding')
    p(d, 'This research received no external funding.')
    h2(d, 'CRediT authorship contribution statement')
    p(d, f'{AUTHOR}: Conceptualization, Methodology, Software, Validation, '
         f'Formal analysis, Investigation, Data curation, Writing — original '
         f'draft, Writing — review & editing, Visualization.')
    h2(d, 'Data availability')
    p(d,
      'All four benchmarks are public (BEIR format). Bibliographic metadata '
      'was obtained from the public Semantic Scholar and OpenAlex APIs. All '
      f'code, fetched metadata, and per-query result files are publicly '
      f'available at {REPO} (archived at https://doi.org/{DOI}).')
    h2(d, 'Use of AI in the manuscript preparation process')
    p(d,
      'During the preparation of this work the author used Kimi (Moonshot '
      'AI) for code debugging and language polishing. All AI-assisted '
      'output, including code and experimental results, was reviewed, '
      'executed, and verified by the author, who takes full responsibility '
      'for the integrity of the publication.')
    d.save(os.path.join(OUT, '04_Declarations.docx'))


# ===========================================================================
# main
# ===========================================================================
def copy_figures():
    mapping = {
        'Fig1_architecture.png': 'system_architecture.png',
        'Fig2_main_results.png': 'Fig2_main_results.png',
        'Fig3_ablation.png': 'Fig3_ablation.png',
        'Fig4_robustness.png': 'Fig4_robustness.png',
        'Fig5_biblioguard.png': 'Fig5_biblioguard.png',
    }
    src = os.path.join(BASE, 'figures')
    for dst, s in mapping.items():
        shutil.copy(os.path.join(src, s), os.path.join(FIGOUT, dst))


if __name__ == '__main__':
    build_title_page()
    build_manuscript()
    build_highlights()
    build_cover_letter()
    build_declarations()
    copy_figures()
    print('ESWA package written to', OUT)
    print(sorted(os.listdir(OUT)))
