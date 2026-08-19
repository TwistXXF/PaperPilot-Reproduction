# -*- coding: utf-8 -*-
"""Audit the ESWA manuscript against results/eswa_tables.json.

Every number the manuscript is supposed to contain must appear in the
extracted docx text. Also runs staleness checks (no leftover IP&M /
two-dataset phrasing). Exit code 0 = all checks pass.
"""
import json
import os
import sys

from docx import Document

import _layout as L

BASE = os.path.dirname(os.path.abspath(__file__))
T = json.load(open(os.path.join(BASE, 'results', 'eswa_tables.json')))
DOCX = L.manuscript_docx()

doc = Document(DOCX)
parts = [p.text for p in doc.paragraphs]
for tbl in doc.tables:
    for row in tbl.rows:
        for cell in row.cells:
            parts.append(cell.text)
TEXT = '\n'.join(parts)

checks = []


def expect(label, value):
    checks.append((label, str(value)))


METHODS = ['BM25', 'LSA-Dense', 'SBERT-Dense', 'BGE-Dense', 'Neural-Hybrid',
           'UMA-RAG', 'LP-RAG', 'CA-HR', 'BGE-Hybrid', 'BGE-CA-HR']
DS = ['scidocs', 'scifact', 'nfcorpus', 'trec-covid']

# 1. main table values
for ds in DS:
    for m in METHODS:
        a = T['main'][ds]['avg'][m]
        expect(f'{ds}/{m} N@10', f"{a['N@10']:.4f}")
        expect(f'{ds}/{m} R@10', f"{a['R@10']:.4f}")

# 2. ablation values
ABL = ['full', '-citation', '-recency', '-dense (alpha=1)',
       '-sparse (alpha=0)', '-rerank (plain hybrid)']
for ds in DS:
    for v in ABL:
        expect(f'ablation {ds} {v}', f"{T['ablation'][ds][v]['N@10']:.4f}")

# 3. robustness quoted values (legacy files contain scalars; newer files
# contain nested metric dictionaries)
for ds, m in [('scifact', 'CA-HR'), ('scifact', 'BM25'),
              ('scifact', 'Neural-Hybrid'),
              ('trec-covid', 'CA-HR'), ('trec-covid', 'BM25'),
              ('trec-covid', 'Neural-Hybrid'),
              ('nfcorpus', 'CA-HR'), ('nfcorpus', 'Neural-Hybrid'),
              ('scidocs', 'BM25'), ('scidocs', 'CA-HR'),
              ('scidocs', 'Neural-Hybrid')]:
    value = T['robust'][ds]['0.4'][m]
    value = value['N@10'] if isinstance(value, dict) else value
    expect(f'robust {ds} 0.4 {m}', f"{value:.4f}")

# 3b. BGE backbone-transfer effect sizes quoted in Section 6.1
for ds in DS:
    d_val = T['main'][ds]['bge_tests']['BGE-CA-HR vs BGE-Hybrid | N@10']['d']
    expect(f'bge transfer d {ds}', f'd = {d_val:.3f}')

# 4. BiblioGuard primary results and mechanism ablation
BG = T['biblioguard']['results']
for ds in DS:
    row = BG[ds]
    expect(f'BiblioGuard fallback {ds}', f"{row['baseline_N@10']:.4f}")
    expect(f'BiblioGuard no-LCB {ds}',
           f"{row['ablation_unconstrained']['N@10']:.4f}")
    expect(f'BiblioGuard guarded {ds}', f"{row['biblioguard_N@10']:.4f}")
    expect(f'BiblioGuard gain {ds}', f"{row['gain_N@10']:+.4f}")
    expect(f'BiblioGuard rates {ds}',
           f"{100*row['ablation_unconstrained']['selection_rate']:.1f}% / "
           f"{100*row['selection_rate']:.1f}%")
expect('BiblioGuard macro gain',
       f"{T['biblioguard']['macro']['gain_N@10']:+.4f}")
for s in ['beta = 0.30 on 141 queries', 'beta = 0.20 on 39',
          'beta = 0.15 on 7', '813 queries return the fallback',
          '96.8%', '67.3%', '45.8%', '54.0%', '-0.00003', '-0.00339',
          '-0.01148']:
    expect(f'BiblioGuard detail {s}', s)

# 5. generation summary (200 paired queries, four datasets)
g = T['generation']
for sysname in ('CA-HR', 'Neural-Hybrid'):
    s = g[sysname]
    expect(f'gen {sysname} relevance mean', f"{s['relevance']['mean']:.2f}")
    expect(f'gen {sysname} faith mean', f"{s['faithfulness']['mean']:.2f}")


def pstr(p):
    return 'p < 0.001' if p < 0.001 else 'p = %.3f' % p


for key in ['paired_relevance', 'paired_faithfulness',
            'paired_citation_precision', 'paired_n_rel_context']:
    expect(f'gen {key} p', pstr(g[key]['wilcoxon_p_two_sided']))
expect('gen paired citprec CA-HR', '%.3f' % g['paired_citation_precision']['mean_CA-HR'])
expect('gen paired citprec NH', '%.3f' % g['paired_citation_precision']['mean_Neural-Hybrid'])
expect('gen rel ctx CA-HR', '%.2f' % g['paired_n_rel_context']['mean_CA-HR'])
expect('gen rel ctx NH', '%.2f' % g['paired_n_rel_context']['mean_Neural-Hybrid'])
# by-dataset citation-precision contrast quoted in text
expect('gen citprec trec-covid', '0.87')
expect('gen citprec scidocs', '0.15')

# 6. dataset facts
for s in ['25,657', '1,000', '5,183', '300', '3,633', '323', '171,332',
          '69.8%', '96.4%', '92.5%', '99.7%', '94.1%', '94.0%']:
    expect(f'dataset fact {s}', s)

# 7. deployment facts
for s in ['10 July 2026', '2 vCPU', '1.6 GB', 'Node.js 20.20.2',
          'MySQL 8.0.46', 'Nginx 1.18', '26 chat conversations',
          '63 messages', '6 registered users']:
    expect(f'deployment {s}', s)

# 8. revised confirmatory family
expect('BiblioGuard scidocs Holm p',
       f"Holm-adjusted p = {BG['scidocs']['wilcoxon_p_holm']:.2g}")
expect('BiblioGuard scidocs d',
       f"paired d = {BG['scidocs']['paired_cohen_d']:.3f}")

# 9. bibliographic diagnostics (Section 5.2, Table 3)
for s in ['0.798', '0.582', '0.498', '0.461', '566', '254', '179',
          '27.4%', '34.7%', '6 vs. 16', 'Mann-Whitney p = 0.55']:
    expect(f'diagnostic {s}', s)

# 10. BGE metadata-weight sensitivity + RRF baselines (Section 6.5, Table 7)
S = T['sensitivity']
for ds in DS:
    expect(f'sens {ds} hybrid', f"{S[ds]['bge_hybrid_N@10']:.4f}")
    expect(f'sens {ds} best', f"{S[ds]['best_N@10']:.4f}")
    for m in ('RRF-MiniLM', 'RRF-BGE'):
        expect(f'rrf {ds} {m}', f"{S[ds]['rrf'][m]['N@10']:.4f}")
expect('sens scidocs best combo', 'beta = 0.3, gamma = 0.0')
expect('sens scidocs significant', 'yes')
for ds in ('scifact', 'nfcorpus', 'trec-covid'):
    assert not S[ds]['any_significant_gain_after_holm'], ds
assert S['scidocs']['any_significant_gain_after_holm']
# Table 4-8 captions and renumbered cross-references
for s in ['Table 4. Cross-fitted BiblioGuard evaluation',
          'Table 5. Retrieval effectiveness',
          'Table 6. CA-HR ablation',
          'Table 7. BGE-backbone metadata-weight sensitivity',
          'Table 8. End-to-end answer quality',
          'Fig. 5. BGE-Hybrid', 'Table 8 shows']:
    expect(f'caption {s}', s)

fails = []
for label, value in checks:
    if value not in TEXT:
        fails.append((label, value))

# staleness checks
stale = []
for bad in ['Information Processing', 'IP&M', 'two benchmarks',
            'two datasets', 'PAV-Agent', 'eight retrieval configurations',
            'Eight retrieval configurations', 'eight methods',
            'eight configurations', '100 paired', '100 test queries',
            'all p > 0.17', 'when—and only when',
            'where metadata coverage is thinner',
            'metadata coverage is thin or',
            'best metadata-aware choice',
            'The stronger BGE-small encoder is the best single',
            'Table 6. Per-query oracle', 'Table 7. End-to-end',
            'Rather than proposing yet another retrieval algorithm',
            'nine atomic citation or recency actions',
            'provides a distribution-free safety guarantee',
            'BiblioGuard has been deployed', 'Table 6 and Fig',
            'Findings of EMNLP, 2022', 'R. Ren, Y. Qu', '679-693',
            'SCIDOCS (computer science; 25,657 documents, '
            '1,000 queries) and SciFact (biomedical;']:
    if bad in TEXT:
        stale.append(bad)

print(f'checks: {len(checks)}, failed: {len(fails)}, stale: {len(stale)}')
for label, value in fails:
    print('MISSING:', label, '->', value)
for s in stale:
    print('STALE:', s)
sys.exit(1 if (fails or stale) else 0)
