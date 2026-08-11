"""verify_paper_numbers.py — audit every number in the manuscript against the
raw data and result files in this repository.

Usage:
    python verify_paper_numbers.py [path/to/01_Manuscript_IPM.docx]

Exit code 0 iff every check passes. Each check prints PASS/FAIL with the
expected (recomputed) and actual (manuscript) value.
"""
import json
import os
import re
import sys

import numpy as np
from scipy import stats as sstats

ROOT = os.path.dirname(os.path.abspath(__file__))
RES = os.path.join(ROOT, 'results')
DATA = os.path.join(ROOT, 'data')

METHODS = ['BM25', 'LSA-Dense', 'SBERT-Dense', 'Neural-Hybrid', 'UMA-RAG',
           'LP-RAG', 'CA-HR']
METRICS = ['R@1', 'R@5', 'R@10', 'N@10', 'MRR']

n_pass = n_fail = 0


def check(name, ok, detail=''):
    global n_pass, n_fail
    if ok:
        n_pass += 1
        print(f'  PASS  {name}')
    else:
        n_fail += 1
        print(f'  FAIL  {name}  {detail}')


def load_pq(ds):
    return np.load(os.path.join(RES, f'{ds}_perquery.npz'), allow_pickle=True)


def mean(pq, m, k):
    return float(np.mean(pq[f'{m}||{k}']))


def cohend(a, b):
    diff = a - b
    sd = diff.std(ddof=1)
    return float(diff.mean() / sd) if sd > 1e-12 else 0.0


def wilcoxon_greater(a, b):
    diff = a - b
    if np.all(np.abs(diff) < 1e-15):
        return 1.0
    try:
        return float(sstats.wilcoxon(a, b, alternative='greater').pvalue)
    except Exception:
        return 1.0


# ---------------------------------------------------------------- data sizes
print('== Dataset sizes ==')
counts = {}
for ds, exp_docs, exp_q in [('scidocs', 25657, 1000), ('scifact', 5183, 300)]:
    corpus = [l for l in open(os.path.join(DATA, ds, 'corpus.jsonl'), encoding='utf8') if l.strip()]
    queries = [l for l in open(os.path.join(DATA, ds, 'queries.jsonl'), encoding='utf8') if l.strip()]
    qrels = [l for l in open(os.path.join(DATA, ds, 'qrels', 'test.tsv'), encoding='utf8') if l.strip()]
    qrels = qrels[1:] if qrels and qrels[0].lower().startswith('query-id') else qrels
    qids = {l.split('\t')[0] if '\t' in l else l.split()[0] for l in qrels}
    counts[ds] = (len(corpus), len(qids))
    check(f'{ds} corpus size = {exp_docs}', len(corpus) == exp_docs, f'got {len(corpus)}')
    check(f'{ds} test queries = {exp_q}', len(qids) == exp_q, f'got {len(qids)}')

print('== Metadata coverage (Semantic Scholar, real API) ==')
for ds, exp_pct in [('scidocs', 99.7), ('scifact', 94.1)]:
    meta = json.load(open(os.path.join(DATA, 'metadata', f'{ds}_metadata.json')))
    matched = sum(1 for v in meta.values() if v is not None)  # API record found
    pct = 100.0 * matched / counts[ds][0]
    check(f'{ds} metadata coverage {pct:.1f}% == {exp_pct}% (paper)',
          round(pct, 1) == exp_pct, f'{matched}/{counts[ds][0]}')

# ------------------------------------------------------------------ tables 1-2
print('== Tables 1-2: main results (recomputed from per-query scores) ==')
import docx  # python-docx

docx_path = sys.argv[1] if len(sys.argv) > 1 else os.path.join(
    ROOT, '..', 'IPM_submission', '01_Manuscript_IPM.docx')
man = docx.Document(docx_path)
pq = {ds: load_pq(ds) for ds in ['scidocs', 'scifact']}

for ti, ds in [(0, 'scidocs'), (1, 'scifact')]:
    tbl = man.tables[ti]
    hdr = [c.text.strip() for c in tbl.rows[0].cells]
    check(f'table {ti} header', hdr == ['Method', 'R@1', 'R@5', 'R@10', 'NDCG@10', 'MRR'], str(hdr))
    colmap = {'R@1': 1, 'R@5': 2, 'R@10': 3, 'N@10': 4, 'MRR': 5}
    for row in tbl.rows[1:]:
        m = row.cells[0].text.strip()
        for k, ci in colmap.items():
            exp = f'{mean(pq[ds], m, k):.4f}'
            got = row.cells[ci].text.strip()
            check(f'{ds} {m} {k}', exp == got, f'paper={got} recomputed={exp}')

# ------------------------------------------------------------------ table 3
print('== Table 3: CA-HR significance (Wilcoxon one-sided, Cohen d, NDCG@10) ==')
tbl = man.tables[2]
for row in tbl.rows[1:]:
    lbl, base = row.cells[0].text.strip(), row.cells[1].text.strip()
    ds = 'scidocs' if lbl == 'SCIDOCS' else 'scifact'
    a = pq[ds]['CA-HR||N@10']
    b = pq[ds][f'{base}||N@10']
    exp = [f'{a.mean():.4f}', f'{b.mean():.4f}',
           f'{wilcoxon_greater(a, b):.4f}', f'{cohend(a, b):.3f}']
    got = [row.cells[i].text.strip() for i in (2, 3, 4, 5)]
    check(f'{lbl} CA-HR vs {base}', exp == got, f'paper={got} recomputed={exp}')

# ------------------------------------------------------------------ table 4
print('== Table 4: CA-HR ablation ==')
abl = {ds: json.load(open(os.path.join(RES, f'{ds}_ablation.json'))) for ds in pq}
tbl = man.tables[3]
for row in tbl.rows[1:]:
    name = row.cells[0].text.strip()
    exp = [f"{abl['scidocs'][name]['N@10']:.4f} / {abl['scidocs'][name]['R@10']:.4f}",
           f"{abl['scifact'][name]['N@10']:.4f} / {abl['scifact'][name]['R@10']:.4f}"]
    got = [row.cells[1].text.strip(), row.cells[2].text.strip()]
    check(f'ablation {name}', exp == got, f'paper={got} recomputed={exp}')

# ------------------------------------------------------------------ table 5
print('== Table 5: robustness under word-drop noise ==')
rob = {ds: json.load(open(os.path.join(RES, f'{ds}_robust.json'))) for ds in pq}
tbl = man.tables[4]
for row in tbl.rows[1:]:
    lev = f"{int(row.cells[0].text.strip().rstrip('%')) / 100:.1f}"
    exp = [f"{rob['scidocs'][lev][m]:.4f}" for m in ['BM25', 'Neural-Hybrid', 'CA-HR']] + \
          [f"{rob['scifact'][lev][m]:.4f}" for m in ['BM25', 'Neural-Hybrid', 'CA-HR']]
    got = [row.cells[i].text.strip() for i in range(1, 7)]
    check(f'noise {lev}', exp == got, f'paper={got} recomputed={exp}')

# ------------------------------------------------------------- in-text claims
print('== In-text numbers ==')
txt = '\n'.join(p.text for p in man.paragraphs)


def in_text(label, value_str):
    check(f'text contains {label}', value_str in txt, f'"{value_str}" not found')


# headline numbers, all recomputed here
rel_gain = (mean(pq['scidocs'], 'SBERT-Dense', 'N@10') /
            mean(pq['scidocs'], 'BM25', 'N@10') - 1) * 100
check('44.7% relative NDCG@10 gain (SCIDOCS SBERT vs BM25)',
      f'{rel_gain:.1f}%' == '44.7%' and '44.7%' in txt, f'{rel_gain:.2f}%')

cit_gain = mean(pq['scidocs'], 'CA-HR', 'N@10') - abl['scidocs']['-citation']['N@10']
check('+0.017 citation-ablation gain on SCIDOCS',
      round(cit_gain, 3) == 0.017 and '+0.017' in txt, f'{cit_gain:.4f}')

rt = {ds: json.load(open(os.path.join(RES, f'{ds}_router.json'))) for ds in pq}
o_s, o_f = rt['scidocs']['oracle']['N@10'], rt['scifact']['oracle']['N@10']
best_s = rt['scidocs']['best_single_N@10']
best_f = mean(pq['scifact'], 'LP-RAG', 'N@10')
check('oracle +0.006 (SCIDOCS)', f'{o_s - best_s:.3f}' == '0.006', f'{o_s - best_s:.4f}')
check('oracle +0.008 (SciFact)', f'{o_f - best_f:.3f}' == '0.008', f'{o_f - best_f:.4f}')
in_text('oracle SCIDOCS 0.2031', f'{o_s:.4f}')
in_text('oracle SciFact 0.7166', f'{o_f:.4f}')
in_text('router acc 76.3%', f"{rt['scidocs']['cv_accuracy'] * 100:.1f}%")
in_text('router acc 91.3%', f"{rt['scifact']['cv_accuracy'] * 100:.1f}%")
in_text("kappa 0.027", f"{rt['scidocs']['kappa']:.3f}")
in_text("kappa -0.005", f"{rt['scifact']['kappa']:.3f}")
in_text('routed N@10 0.1896', f"{rt['scidocs']['routed_system']['N@10']:.4f}")
in_text('routed N@10 0.7042', f"{rt['scifact']['routed_system']['N@10']:.4f}")

for s in ['0.2164', '0.3707', '0.7081', '0.7080', '0.6496', '0.6451',
          '0.0867', '0.4260', '0.2022', '0.1974', '-0.187',
          '0.1860', '0.1365', '0.1578', '0.1208', '0.6923', '0.6697',
          '99.7%', '94.1%', '25,657', '5,183']:
    in_text(s, s)

# cross-check a few against data
check('text 0.2164 == SBERT-Dense N@10', f"{mean(pq['scidocs'], 'SBERT-Dense', 'N@10'):.4f}" == '0.2164')
check('text 0.2022 == -recency ablation', f"{abl['scidocs']['-recency']['N@10']:.4f}" == '0.2022')
d_sbert = cohend(pq['scidocs']['CA-HR||N@10'], pq['scidocs']['SBERT-Dense||N@10'])
check('text d=-0.187 (CA-HR vs SBERT, SCIDOCS)', f'{d_sbert:.3f}' == '-0.187', f'{d_sbert:.3f}')
for v, lev, m, ds in [('0.1860', '0.1', 'CA-HR', 'scidocs'), ('0.1365', '0.1', 'BM25', 'scidocs'),
                      ('0.1578', '0.3', 'CA-HR', 'scidocs'), ('0.1208', '0.3', 'BM25', 'scidocs'),
                      ('0.6923', '0.4', 'CA-HR', 'scifact'), ('0.6697', '0.4', 'BM25', 'scifact')]:
    check(f'robustness {ds} {m} @{lev} == {v}', f"{rob[ds][lev][m]:.4f}" == v)

print(f'\n{n_pass} passed, {n_fail} failed')
sys.exit(1 if n_fail else 0)
